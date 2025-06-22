import sys
import os
import json
import logging
import requests
import base64
import io
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QLineEdit, QTabWidget,
    QMessageBox, QTextEdit, QGroupBox, QFormLayout, QInputDialog,
    QListWidget, QListWidgetItem, QDialog, QComboBox
)
from PyQt5.QtCore import Qt
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from Crypto.Protocol.KDF import PBKDF2
from dilithium_py.ml_dsa import ML_DSA_44
from PyPDF2 import PdfReader, PdfWriter
from PyPDF2.generic import NameObject, TextStringObject
from datetime import datetime, timedelta
import qrcode
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import tempfile
import hashlib
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cấu hình logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Unique identifier for our metadata
METADATA_IDENTIFIER = "MLDSA_METADATA_JSON:"

def verify_certificate(certificate):
    """Verify a certificate using CA's public key"""
    try:
        # Read CA info from file
        ca_info_path = os.path.join(os.path.expanduser("~"), ".pdf_verifier", "keys", "ca_info.json")
        if not os.path.exists(ca_info_path):
            print("CA info file not found")
            return False
            
        with open(ca_info_path, 'r') as f:
            ca_info = json.load(f)
        
        # Use the verified CA public key
        ca_public_key = base64.b64decode(ca_info['ca_public_key'])
        
        # Extract payload and signature from certificate
        payload = certificate['payload']
        signature = base64.b64decode(certificate['signature'])
        
        # Verify signature
        return ML_DSA_44.verify(ca_public_key, json.dumps(payload, sort_keys=True).encode(), signature)
    except Exception as e:
        print(f"Error verifying certificate: {e}")
        return False

class PDFVerifier(QMainWindow):
    def __init__(self):
        super().__init__()
        self.API_BASE_URL = "http://localhost:8000"
        self.current_user = None
        self.current_user_role = None
        self.keys_dir = os.path.join(os.path.expanduser("~"), ".pdf_verifier", "keys")
        os.makedirs(self.keys_dir, exist_ok=True)
        self.init_ui()

    def _store_metadata_in_docx(self, doc, metadata):
        """Stores metadata as a hidden paragraph at the end of the DOCX."""
        # Remove old metadata paragraph if it exists
        self._remove_metadata_paragraph(doc)
            
        metadata_str = json.dumps(metadata)
        full_text = f"{METADATA_IDENTIFIER}{metadata_str}"
    
        # Add new paragraph with the metadata
        p = doc.add_paragraph()
        run = p.add_run(full_text)
    
        # Make it "hidden"
        font = run.font
        font.size = Pt(1)
        font.color.rgb = RGBColor(255, 255, 255) # White color

    def _extract_metadata_from_docx(self, doc):
        """Extracts metadata from a hidden paragraph in the DOCX."""
        print(f"DEBUG: Bắt đầu trích xuất metadata từ {len(doc.paragraphs)} paragraphs")
        
        for i, p in enumerate(reversed(doc.paragraphs)):
            print(f"DEBUG: Paragraph {len(doc.paragraphs) - i - 1}: '{p.text.strip()[:50]}...'")
            if p.text.strip().startswith(METADATA_IDENTIFIER):
                json_str = p.text.strip()[len(METADATA_IDENTIFIER):]
                print(f"DEBUG: Tìm thấy metadata identifier, JSON string: {json_str[:100]}...")
                try:
                    metadata = json.loads(json_str)
                    print(f"DEBUG: Parse JSON thành công: {metadata}")
                    return metadata
                except json.JSONDecodeError as e:
                    print(f"DEBUG: Lỗi parse JSON: {e}")
                    return None
        print(f"DEBUG: Không tìm thấy metadata identifier trong document")
        return None

    def _remove_metadata_paragraph(self, doc):
        """Finds and removes the metadata paragraph from the document."""
        for p in reversed(doc.paragraphs):
            if p.text.strip().startswith(METADATA_IDENTIFIER):
                p_element = p._element
                p_element.getparent().remove(p_element)
                return True # Found and removed
        return False # Not found

    def init_ui(self):
        """Khởi tạo giao diện"""
        self.setWindowTitle('Hệ thống Ký và Xác thực PDF')
        self.setGeometry(100, 100, 1000, 600)

        # Widget chính
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        self.main_layout = QVBoxLayout(main_widget)

        # Tạo login widget
        self.login_widget = QWidget()
        self.init_login_ui()
        self.main_layout.addWidget(self.login_widget)

        # Tạo main content widget (sẽ ẩn ban đầu)
        self.content_widget = QWidget()
        self.init_content_ui()
        self.content_widget.hide()
        self.main_layout.addWidget(self.content_widget)

    def init_login_ui(self):
        """Khởi tạo giao diện đăng nhập"""
        login_layout = QVBoxLayout(self.login_widget)
        
        # Tạo group box cho phần đăng nhập
        login_group = QGroupBox("Đăng nhập")
        login_form = QFormLayout()
        
        self.username_input = QLineEdit()
        login_form.addRow("Username:", self.username_input)
        
        # Thêm layout ngang cho các nút
        button_layout = QHBoxLayout()
        
        login_btn = QPushButton("Đăng nhập")
        login_btn.clicked.connect(self.login)
        button_layout.addWidget(login_btn)
        
        register_btn = QPushButton("Đăng ký")
        register_btn.clicked.connect(self.show_register_dialog)
        button_layout.addWidget(register_btn)
        
        login_form.addRow("", button_layout)
        login_group.setLayout(login_form)
        
        # Thêm vào layout chính
        login_layout.addWidget(login_group)
        login_layout.addStretch()

    def init_content_ui(self):
        """Khởi tạo giao diện chính sau khi đăng nhập"""
        content_layout = QVBoxLayout(self.content_widget)
        
        # Thêm header với thông tin người dùng và nút đăng xuất
        header_layout = QHBoxLayout()
        self.user_label = QLabel()
        header_layout.addWidget(self.user_label)
        
        logout_btn = QPushButton("Đăng xuất")
        logout_btn.clicked.connect(self.logout)
        header_layout.addWidget(logout_btn)
        
        content_layout.addLayout(header_layout)
        
        # Tab widget
        self.tabs = QTabWidget()
        
        # Tab Bên bán (khởi tạo nhưng chưa thêm vào tabs)
        self.seller_tab = QWidget()
        seller_layout = QVBoxLayout()
        
        # Product management section
        product_group = QGroupBox("Quản lý sản phẩm")
        product_layout = QVBoxLayout()
        
        # Product input fields
        product_form = QFormLayout()
        self.product_name_input = QLineEdit()
        self.product_price_input = QLineEdit()
        self.product_price_input.setPlaceholderText("Ví dụ: 100000")
        product_form.addRow("Tên sản phẩm:", self.product_name_input)
        product_form.addRow("Giá (VNĐ):", self.product_price_input)
        product_layout.addLayout(product_form)
        
        # Add product button
        add_product_btn = QPushButton("Thêm sản phẩm")
        add_product_btn.clicked.connect(self.add_product)
        product_layout.addWidget(add_product_btn)
        
        # My products list
        self.my_products_list = QListWidget()
        product_layout.addWidget(QLabel("Sản phẩm của tôi:"))
        product_layout.addWidget(self.my_products_list)
        
        refresh_products_btn = QPushButton("Làm mới danh sách sản phẩm")
        refresh_products_btn.clicked.connect(self.refresh_my_products)
        product_layout.addWidget(refresh_products_btn)
        
        product_group.setLayout(product_layout)
        seller_layout.addWidget(product_group)
        
        # Thêm danh sách order từ bên mua
        order_group = QGroupBox("Danh sách order từ bên mua")
        order_layout = QVBoxLayout()
        self.order_list = QListWidget()
        self.order_list.itemDoubleClicked.connect(self.download_and_verify_order)
        order_layout.addWidget(self.order_list)
        refresh_order_btn = QPushButton("Làm mới danh sách order")
        refresh_order_btn.clicked.connect(self.refresh_order_list)
        order_layout.addWidget(refresh_order_btn)
        order_group.setLayout(order_layout)
        seller_layout.addWidget(order_group)
        
        # Thêm danh sách hóa đơn của tôi
        my_invoice_group = QGroupBox("Hóa đơn của tôi")
        my_invoice_layout = QVBoxLayout()
        self.my_invoice_list = QListWidget()
        self.my_invoice_list.itemDoubleClicked.connect(self.download_and_verify_invoice)
        my_invoice_layout.addWidget(self.my_invoice_list)
        refresh_my_invoice_btn = QPushButton("Làm mới danh sách hóa đơn")
        refresh_my_invoice_btn.clicked.connect(self.refresh_my_invoice_list)
        my_invoice_layout.addWidget(refresh_my_invoice_btn)
        my_invoice_group.setLayout(my_invoice_layout)
        seller_layout.addWidget(my_invoice_group)
        
        self.seller_tab.setLayout(seller_layout)

        # Tab Bên mua (khởi tạo nhưng chưa thêm vào tabs)
        self.buyer_tab = QWidget()
        buyer_layout = QVBoxLayout()
        
        # Shopping section
        shopping_group = QGroupBox("Mua sắm")
        shopping_layout = QVBoxLayout()
        
        # Products list
        self.products_list = QListWidget()
        self.products_list.itemDoubleClicked.connect(self.select_product_for_order)
        shopping_layout.addWidget(QLabel("Danh sách sản phẩm:"))
        shopping_layout.addWidget(self.products_list)
        
        refresh_products_btn = QPushButton("Làm mới danh sách sản phẩm")
        refresh_products_btn.clicked.connect(self.refresh_products)
        shopping_layout.addWidget(refresh_products_btn)
        
        # Shopping cart
        cart_group = QGroupBox("Giỏ hàng")
        cart_layout = QVBoxLayout()
        
        self.cart_list = QListWidget()
        cart_layout.addWidget(self.cart_list)
        
        # Total price
        self.total_price_label = QLabel("Tổng tiền: 0 VNĐ")
        cart_layout.addWidget(self.total_price_label)
        
        # Cart buttons
        cart_buttons_layout = QHBoxLayout()
        clear_cart_btn = QPushButton("Xóa giỏ hàng")
        clear_cart_btn.clicked.connect(self.clear_cart)
        create_order_btn = QPushButton("Tạo đơn hàng")
        create_order_btn.clicked.connect(self.create_order_from_cart)
        cart_buttons_layout.addWidget(clear_cart_btn)
        cart_buttons_layout.addWidget(create_order_btn)
        cart_layout.addLayout(cart_buttons_layout)
        
        cart_group.setLayout(cart_layout)
        shopping_layout.addWidget(cart_group)
        
        shopping_group.setLayout(shopping_layout)
        buyer_layout.addWidget(shopping_group)
        
        # Invoice list section
        invoice_group = QGroupBox("Danh sách hóa đơn")
        invoice_layout = QVBoxLayout()
        
        self.invoice_list = QListWidget()
        self.invoice_list.itemDoubleClicked.connect(self.download_and_verify_invoice)
        invoice_layout.addWidget(self.invoice_list)
        
        refresh_btn = QPushButton("Làm mới danh sách")
        refresh_btn.clicked.connect(self.refresh_invoice_list)
        invoice_layout.addWidget(refresh_btn)
        
        invoice_group.setLayout(invoice_layout)
        buyer_layout.addWidget(invoice_group)
        
        # Thêm danh sách order đã gửi
        my_order_group = QGroupBox("Order đã gửi")
        my_order_layout = QVBoxLayout()
        self.my_order_list = QListWidget()
        my_order_layout.addWidget(self.my_order_list)
        refresh_my_order_btn = QPushButton("Làm mới danh sách order")
        refresh_my_order_btn.clicked.connect(self.refresh_my_order_list)
        my_order_layout.addWidget(refresh_my_order_btn)
        my_order_group.setLayout(my_order_layout)
        buyer_layout.addWidget(my_order_group)
        
        self.buyer_tab.setLayout(buyer_layout)

        content_layout.addWidget(self.tabs)

    def login(self):
        """Đăng nhập với username"""
        username = self.username_input.text().strip()
        if not username:
            QMessageBox.warning(self, "Lỗi", "Vui lòng nhập username")
            return
            
        try:
            # Kiểm tra user tồn tại và lấy vai trò
            response = requests.get(
                f"{self.API_BASE_URL}/check-user",
                params={"username": username}
            )
            response_data = response.json()
            
            if response.status_code == 200 and response_data.get('message') == 'User exists':
                self.current_user = username
                self.current_user_role = response_data.get('role')
                self.user_label.setText(f"Đã đăng nhập: {username} ({self.current_user_role})")
                
                # Ẩn login widget và hiện content widget
                self.login_widget.hide()
                self.content_widget.show()
                
                self.update_ui_for_role()
                
                QMessageBox.information(self, "Thành công", f"Đăng nhập thành công: {username}")
            else:
                QMessageBox.warning(self, "Lỗi", response_data.get('message', 'Username không tồn tại hoặc lỗi server'))
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi đăng nhập: {str(e)}")

    def logout(self):
        """Đăng xuất"""
        self.current_user = None
        self.current_user_role = None
        self.username_input.clear()
        
        # Xóa tất cả các tab khi đăng xuất
        self.tabs.clear()

        # Ẩn content widget và hiện login widget
        self.content_widget.hide()
        self.login_widget.show()
        
        QMessageBox.information(self, "Thông báo", "Đã đăng xuất thành công")

    def update_ui_for_role(self):
        """Cập nhật giao diện dựa trên vai trò của người dùng"""
        # Xóa tất cả các tab hiện có
        self.tabs.clear() 

        if self.current_user_role == 'seller':
            self.tabs.addTab(self.seller_tab, "Bên bán")
            self.tabs.setCurrentWidget(self.seller_tab)
            self.refresh_my_products()
            self.refresh_order_list()
            self.refresh_my_invoice_list()
        elif self.current_user_role == 'buyer':
            self.tabs.addTab(self.buyer_tab, "Bên mua")
            self.tabs.setCurrentWidget(self.buyer_tab)
            self.refresh_products()
            self.refresh_invoice_list()
            self.refresh_my_order_list()
        else:
            QMessageBox.warning(self, "Lỗi", "Vai trò người dùng không xác định. Vui lòng liên hệ quản trị viên.")

    def select_sign_file(self):
        """Chọn file DOCX để ký"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Chọn file DOCX", "", "Word Documents (*.docx)"
        )
        if file_path:
            self.sign_path.setText(file_path)

    def sign_and_upload_docx(self):
        """Ký và upload file DOCX"""
        if not self.current_user:
            QMessageBox.warning(self, "Lỗi", "Vui lòng đăng nhập trước")
            return
        file_path = self.sign_path.text()
        if not file_path:
            QMessageBox.warning(self, "Lỗi", "Vui lòng chọn file DOCX")
            return
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu:", QLineEdit.Password
        )
        if not ok or not password:
            return
        try:
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key")
                return
            
            doc = Document(file_path)
            timestamp = datetime.now().isoformat()

            # Bước 1: Thêm QR code vào document
            qr_image = self.create_qr_code(self.current_user, timestamp)
            self.add_qr_to_docx(doc, qr_image)

            # Bước 2: Thêm metadata TRƯỚC khi ký
            metadata = {
                'signer': self.current_user,
                'sign_time': timestamp
            }
            self._store_metadata_in_docx(doc, metadata)

            # Bước 3: Lấy bytes của document (với QR và metadata) để ký
            buffer_to_sign = io.BytesIO()
            doc.save(buffer_to_sign)
            bytes_to_sign = buffer_to_sign.getvalue()

            # Bước 4: Ký
            signature = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64 = base64.b64encode(signature).decode('utf-8')

            # Bước 5: Cập nhật metadata với chữ ký thật và bytes đã ký
            metadata['signature'] = signature_b64
            metadata['signed_data_bytes'] = base64.b64encode(bytes_to_sign).decode('utf-8')
            metadata['signed_data_hash'] = hashlib.sha256(bytes_to_sign).hexdigest()
            self._store_metadata_in_docx(doc, metadata)
            
            # Bước 6: Lưu file DOCX cuối cùng
            final_docx_path_for_upload = file_path.replace('.docx', '_signed_with_qr.docx')
            doc.save(final_docx_path_for_upload)
            
            # Bước 7: Upload file đã ký và có QR lên server
            with open(final_docx_path_for_upload, 'rb') as f:
                files = {'pdf': f}
                data = {
                    'signer_name': self.current_user,
                    'signature': signature_b64,
                    'timestamp': timestamp
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-invoice/",
                    files=files,
                    data=data
                )
                response.raise_for_status()
            
            QMessageBox.information(
                self, "Thành công",
                f"Ký và upload DOCX thành công. File đã ký được lưu tại:\n{final_docx_path_for_upload}"
            )
            self.sign_path.clear()
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi ký và upload DOCX: {str(e)}")

    def refresh_invoice_list(self):
        """Làm mới danh sách hóa đơn"""
        if not self.current_user:
            return

        try:
            response = requests.get(
                f"{self.API_BASE_URL}/list-invoices/"
            )
            response.raise_for_status()
            
            invoices = response.json()
            self.invoice_list.clear()
            
            for invoice in invoices:
                item = QListWidgetItem(
                    f"Hóa đơn từ {invoice['signer_name']} - {invoice['timestamp']}"
                )
                item.setData(Qt.UserRole, invoice)
                self.invoice_list.addItem(item)
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách hóa đơn: {str(e)}")

    def download_and_verify_invoice(self, item):
        """Tải và xác thực hóa đơn DOCX"""
        invoice = item.data(Qt.UserRole)
        try:
            # Tải invoice
            response = requests.get(
                f"{self.API_BASE_URL}/download-invoice/{invoice['id']}"
            )
            response.raise_for_status()
            
            temp_path = os.path.join(os.path.expanduser("~"), "Downloads", f"invoice_{invoice['id']}_from_{invoice['signer_name']}.docx")
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            
            # Đọc file DOCX đã ký
            doc = Document(temp_path)
            
            # Bước 1: Trích xuất metadata và chữ ký
            metadata = self._extract_metadata_from_docx(doc)
            print(f"DEBUG: Metadata trích xuất được: {metadata}")
            
            if not metadata or 'signature' not in metadata:
                QMessageBox.warning(self, "Lỗi", "File DOCX không có metadata hoặc chữ ký.")
                return

            signature_b64_seller = metadata.get('signature')
            signature_seller = base64.b64decode(signature_b64_seller)
            signer_seller = metadata.get('signer')
            sign_time_seller = metadata.get('sign_time')
            stored_hash = metadata.get('signed_data_hash')
            stored_bytes_b64 = metadata.get('signed_data_bytes')

            # Lấy certificate của người bán
            response_cert = requests.get(
                f"{self.API_BASE_URL}/get-certificate/",
                params={'username': signer_seller}
            )
            response_cert.raise_for_status()
            seller_certificate = response_cert.json()['certificate']

            # Xác minh certificate
            if not verify_certificate(seller_certificate):
                QMessageBox.warning(self, "Lỗi", "Certificate của người bán không hợp lệ")
                return

            # Lấy public key từ certificate đã xác minh
            public_key_seller = base64.b64decode(seller_certificate['payload']['public_key'])
            
            # Bước 2: Sử dụng bytes đã lưu trực tiếp thay vì tạo lại document
            if stored_bytes_b64:
                docx_bytes_to_verify = base64.b64decode(stored_bytes_b64)
                print(f"DEBUG: Sử dụng bytes đã lưu - length: {len(docx_bytes_to_verify)} bytes")
                print(f"DEBUG: Hash của bytes đã lưu: {hashlib.sha256(docx_bytes_to_verify).hexdigest()[:16]}")
            else:
                # Fallback: Tạo lại document nếu không có bytes đã lưu (cho các file cũ)
                print(f"DEBUG: Không có bytes đã lưu, tạo lại document...")
                buf_seller = io.BytesIO()
                doc.save(buf_seller)
                docx_bytes_to_verify = buf_seller.getvalue()
                print(f"DEBUG: Tạo lại document - length: {len(docx_bytes_to_verify)} bytes")
            
            # Bước 3: Xác minh chữ ký
            try:
                is_valid_seller_signature = ML_DSA_44.verify(public_key_seller, docx_bytes_to_verify, signature_seller)
                print(f"DEBUG: Kết quả xác thực: {is_valid_seller_signature}")
            except Exception as verify_error:
                print(f"DEBUG: Lỗi khi xác thực: {str(verify_error)}")
                is_valid_seller_signature = False

            # Hiển thị kết quả xác thực
            result_text = f"Kết quả xác thực hóa đơn:\n\n"
            result_text += f"Người bán: {signer_seller}\n"
            result_text += f"Thời gian ký: {sign_time_seller}\n"
            result_text += f"Chữ ký người bán: {'Hợp lệ' if is_valid_seller_signature else 'Không hợp lệ'}\n\n"
            
            if is_valid_seller_signature:
                result_text += "✅ Hóa đơn đã được xác thực thành công!"
            else:
                result_text += "❌ Hóa đơn không hợp lệ!"

            QMessageBox.information(self, "Kết quả xác thực", result_text)
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi xác thực hóa đơn: {str(e)}")

    def get_key_paths(self, username=None):
        """Lấy đường dẫn đến file khóa"""
        # Sử dụng username được truyền vào hoặc current_user
        target_username = username or self.current_user
        if not target_username:
            return None, None
        
        private_key_path = os.path.join(self.keys_dir, f"{target_username}.private.pem")
        public_key_path = os.path.join(self.keys_dir, f"{target_username}.public.pem")
        return private_key_path, public_key_path

    def save_keys(self, public_key, private_key, password, username=None):
        """Lưu cặp khóa vào file"""
        try:
            # Đảm bảo thư mục keys tồn tại
            os.makedirs(self.keys_dir, exist_ok=True)
            
            # Sử dụng username được truyền vào hoặc current_user
            target_username = username or self.current_user
            if not target_username:
                print("Không có username để lưu khóa")  # Debug log
                return False
            
            private_key_path = os.path.join(self.keys_dir, f"{target_username}.private.pem")
            public_key_path = os.path.join(self.keys_dir, f"{target_username}.public.pem")
            
            print(f"Đang lưu khóa vào: {private_key_path} và {public_key_path}")  # Debug log

            # Lưu public key
            with open(public_key_path, 'w') as f:
                json.dump({
                    'public_key': base64.b64encode(public_key).decode('utf-8')
                }, f)

            # Mã hóa và lưu private key
            salt = get_random_bytes(16)
            iv = get_random_bytes(16)
            key = PBKDF2(password.encode(), salt, dkLen=32, count=10000)
            cipher = AES.new(key, AES.MODE_CBC, iv)
            
            # Pad private key
            padded_sk = private_key + b'\0' * (16 - len(private_key) % 16)
            encrypted_sk = cipher.encrypt(padded_sk)

            with open(private_key_path, 'w') as f:
                json.dump({
                    'salt': base64.b64encode(salt).decode('utf-8'),
                    'iv': base64.b64encode(iv).decode('utf-8'),
                    'encrypted_key': base64.b64encode(encrypted_sk).decode('utf-8')
                }, f)

            return True
        except Exception as e:
            logging.error(f"Lỗi khi lưu khóa: {str(e)}")
            print(f"Chi tiết lỗi: {str(e)}")  # Debug log
            return False

    def load_private_key(self, password):
        """Đọc và giải mã private key"""
        try:
            private_key_path, _ = self.get_key_paths()
            if not private_key_path or not os.path.exists(private_key_path):
                return None

            with open(private_key_path, 'r') as f:
                key_data = json.load(f)
                salt = base64.b64decode(key_data['salt'])
                iv = base64.b64decode(key_data['iv'])
                encrypted_key = base64.b64decode(key_data['encrypted_key'])

            key = PBKDF2(password.encode(), salt, dkLen=32, count=10000)
            cipher = AES.new(key, AES.MODE_CBC, iv)
            decrypted_key = cipher.decrypt(encrypted_key)
            return decrypted_key.rstrip(b'\0')

        except Exception as e:
            logging.error(f"Lỗi khi đọc private key: {str(e)}")
            return None

    def select_order_sign_file(self):
        """Chọn file Order DOCX để ký"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Chọn file Order DOCX", "", "Word Documents (*.docx)"
        )
        if file_path:
            self.order_sign_path.setText(file_path)

    def sign_and_upload_order_docx(self):
        """Bên mua ký và upload Order DOCX"""
        if not self.current_user:
            QMessageBox.warning(self, "Lỗi", "Vui lòng đăng nhập trước")
            return
        file_path = self.order_sign_path.text()
        if not file_path:
            QMessageBox.warning(self, "Lỗi", "Vui lòng chọn file Order DOCX")
            return
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu:", QLineEdit.Password
        )
        if not ok or not password:
            return
        try:
            print(f"DEBUG: Bắt đầu ký file {file_path}")
            
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key")
                return
            
            doc = Document(file_path)
            timestamp = datetime.now().isoformat()

            # Bước 1: Thêm QR code
            qr_image = self.create_qr_code(self.current_user, timestamp)
            self.add_qr_to_docx(doc, qr_image)
            
            # Bước 2: Thêm metadata TRƯỚC khi ký (không có chữ ký)
            metadata = {
                'buyer': self.current_user,
                'sign_time': timestamp
            }
            print(f"DEBUG: Metadata để ký: {json.dumps(metadata, sort_keys=True)}")
            self._store_metadata_in_docx(doc, metadata)
            
            # Bước 3: Lấy bytes để ký (bao gồm cả metadata)
            buffer_to_sign = io.BytesIO()
            doc.save(buffer_to_sign)
            bytes_to_sign = buffer_to_sign.getvalue()
            
            print(f"DEBUG: Bytes để ký - length: {len(bytes_to_sign)} bytes")
            print(f"DEBUG: Hash của bytes để ký: {hashlib.sha256(bytes_to_sign).hexdigest()[:16]}")

            # Bước 4: Ký
            signature = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64 = base64.b64encode(signature).decode('utf-8')
            
            print(f"DEBUG: Chữ ký tạo ra - length: {len(signature)} bytes")
            
            # Bước 5: Tạo metadata cuối cùng với chữ ký và hash của bytes đã ký
            final_metadata = {
                'buyer': self.current_user,
                'signature': signature_b64,
                'sign_time': timestamp,
                'signed_data_hash': hashlib.sha256(bytes_to_sign).hexdigest(),
                'signed_data_bytes': base64.b64encode(bytes_to_sign).decode('utf-8')
            }
            
            # Bước 6: Tạo file cuối cùng với metadata đầy đủ
            final_doc = Document(file_path)
            qr_image_final = self.create_qr_code(self.current_user, timestamp)
            self.add_qr_to_docx(final_doc, qr_image_final)
            self._store_metadata_in_docx(final_doc, final_metadata)
            
            final_docx_path_for_upload = file_path.replace('.docx', '_signed_with_qr.docx')
            final_doc.save(final_docx_path_for_upload)

            # Bước 7: Upload file đã ký
            with open(final_docx_path_for_upload, 'rb') as f:
                files = {'pdf': f}
                data = {
                    'buyer_name': self.current_user,
                    'signature': signature_b64,
                    'timestamp': timestamp
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-order/",
                    files=files,
                    data=data
                )
                response.raise_for_status()
            
            print(f"DEBUG: Upload thành công, order ID: {response.json().get('id', 'unknown')}")
            
            QMessageBox.information(
                self, "Thành công",
                f"Ký và upload Order DOCX thành công. File đã ký được lưu tại:\n{final_docx_path_for_upload}"
            )
            self.order_sign_path.clear()
            
        except Exception as e:
            print(f"DEBUG: Exception trong sign_and_upload_order_docx: {str(e)}")
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi ký và upload Order DOCX: {str(e)}")

    def refresh_order_list(self):
        try:
            response = requests.get(f"{self.API_BASE_URL}/list-orders/")
            response.raise_for_status()
            orders = response.json()
            self.order_list.clear()
            for order in orders:
                item = QListWidgetItem(
                    f"Order từ {order['buyer_name']} - {order['timestamp']}"
                )
                item.setData(Qt.UserRole, order)
                self.order_list.addItem(item)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách order: {str(e)}")

    def refresh_my_order_list(self):
        try:
            response = requests.get(f"{self.API_BASE_URL}/list-orders/")
            response.raise_for_status()
            orders = response.json()
            self.my_order_list.clear()
            for order in orders:
                if order['buyer_name'] == self.current_user:
                    item = QListWidgetItem(
                        f"Order đã gửi - {order['timestamp']}"
                    )
                    item.setData(Qt.UserRole, order)
                    self.my_order_list.addItem(item)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách order đã gửi: {str(e)}")

    def create_qr_code(self, signer_name, timestamp):
        """Tạo mã QR chứa thông tin người ký và thời gian"""
        # Tạo dữ liệu cho mã QR
        qr_data = {
            "signer": signer_name,
            "timestamp": timestamp,
        }
        qr_data_str = json.dumps(qr_data)
        
        # Tạo mã QR
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(qr_data_str)
        qr.make(fit=True)
        
        # Tạo hình ảnh mã QR
        qr_image = qr.make_image(fill_color="black", back_color="white")
        
        # Chuyển đổi hình ảnh thành bytes
        img_byte_arr = BytesIO()
        qr_image.save(img_byte_arr, format='PNG')
        return img_byte_arr.getvalue()

    def add_qr_to_docx(self, doc, qr_image_bytes):
        """Thêm mã QR vào cuối file Word DOCX, modifying the object in place."""
        try:
            # Thêm một trang mới (section break) để chứa QR code
            doc.add_page_break()
            
            # Thêm tiêu đề cho phần QR code
            qr_heading = doc.add_heading('Digital Signature QR Code', level=1)
            qr_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm đoạn văn bản mô tả
            description = doc.add_paragraph('This QR code contains the digital signature information for this document.')
            description.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Lưu QR code vào file tạm thời
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_qr_file:
                temp_qr_file.write(qr_image_bytes)
                temp_qr_path = temp_qr_file.name
            
            # Thêm QR code vào document
            qr_paragraph = doc.add_paragraph()
            qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr_run = qr_paragraph.add_run()
            qr_run.add_picture(temp_qr_path, width=Inches(2.5), height=Inches(2.5))
            
            # Xóa file QR tạm thời
            os.remove(temp_qr_path)
            
            # Thêm thông tin bổ sung
            info_paragraph = doc.add_paragraph()
            info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_paragraph.add_run('Scan this QR code to verify the digital signature of this document.')
            
        except Exception as e:
            print(f"Error adding QR to DOCX: {str(e)}")
            raise

    def download_and_verify_order(self, item):
        """Tải và xác thực order DOCX"""
        order = item.data(Qt.UserRole)
        try:
            print(f"DEBUG: Bắt đầu xác thực order {order['id']} từ {order['buyer_name']}")
            
            response = requests.get(
                f"{self.API_BASE_URL}/download-order/{order['id']}"
            )
            response.raise_for_status()
            
            temp_path = os.path.join(os.path.expanduser("~"), "Downloads", f"order_{order['id']}_from_{order['buyer_name']}.docx")
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            
            print(f"DEBUG: Đã tải order về {temp_path}, kích thước: {len(response.content)} bytes")
            
            # Đọc file DOCX đã ký
            doc = Document(temp_path)
            
            # Bước 1: Trích xuất metadata và chữ ký
            metadata = self._extract_metadata_from_docx(doc)
            print(f"DEBUG: Metadata trích xuất được: {metadata}")
            
            if not metadata or 'signature' not in metadata:
                QMessageBox.warning(self, "Lỗi", "File DOCX không có metadata hoặc chữ ký.")
                return
            
            signature_b64 = metadata.get('signature')
            signature = base64.b64decode(signature_b64)
            buyer = metadata.get('buyer')
            sign_time = metadata.get('sign_time')
            stored_hash = metadata.get('signed_data_hash')
            stored_bytes_b64 = metadata.get('signed_data_bytes')
            
            print(f"DEBUG: Chữ ký từ metadata - length: {len(signature)} bytes")
            print(f"DEBUG: Người mua: {buyer}, thời gian ký: {sign_time}")
            print(f"DEBUG: Hash đã lưu trong metadata: {stored_hash}")
            
            if not signature_b64 or not buyer:
                QMessageBox.warning(self, "Lỗi", "File DOCX không có chữ ký hợp lệ")
                return
            
            # Lấy certificate của người mua
            response_cert = requests.get(
                f"{self.API_BASE_URL}/get-certificate/",
                params={'username': buyer}
            )
            response_cert.raise_for_status()
            buyer_certificate = response_cert.json()['certificate']

            # Xác minh certificate
            if not verify_certificate(buyer_certificate):
                QMessageBox.warning(self, "Lỗi", "Certificate của người mua không hợp lệ")
                return

            # Lấy public key từ certificate đã xác minh
            public_key_buyer = base64.b64decode(buyer_certificate['payload']['public_key'])
            print(f"DEBUG: Public key length: {len(public_key_buyer)} bytes")
            
            # Bước 2: Sử dụng bytes đã lưu trực tiếp thay vì tạo lại document
            if stored_bytes_b64:
                docx_bytes_to_verify = base64.b64decode(stored_bytes_b64)
                print(f"DEBUG: Sử dụng bytes đã lưu - length: {len(docx_bytes_to_verify)} bytes")
                print(f"DEBUG: Hash của bytes đã lưu: {hashlib.sha256(docx_bytes_to_verify).hexdigest()[:16]}")
            else:
                # Fallback: Tạo lại document nếu không có bytes đã lưu
                print(f"DEBUG: Không có bytes đã lưu, tạo lại document...")
                QMessageBox.warning(self, "Lỗi", "File không có dữ liệu đã ký. Không thể xác thực.")
                return
            
            # Bước 3: Xác minh chữ ký
            try:
                is_valid_signature = ML_DSA_44.verify(public_key_buyer, docx_bytes_to_verify, signature)
                print(f"DEBUG: Kết quả xác thực: {is_valid_signature}")
            except Exception as verify_error:
                print(f"DEBUG: Lỗi khi xác thực: {str(verify_error)}")
                is_valid_signature = False
            
            if is_valid_signature:
                reply = QMessageBox.information(
                    self, "Xác thực Order thành công!",
                    f"Order từ {buyer} - Thời gian: {sign_time}\n"
                    f"Chữ ký của người mua hợp lệ.\n\nBạn có muốn ký và tạo Hóa đơn (Invoice) từ Order này không?",
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
                )
                if reply == QMessageBox.Yes:
                    self.sign_downloaded_order_as_invoice(temp_path, buyer)
                elif reply == QMessageBox.No:
                    QMessageBox.information(self, "Thông báo", "Order đã được tải về và xác thực.")
            else:
                # Hiển thị thông tin debug chi tiết hơn
                debug_info = f"DEBUG INFO:\n"
                debug_info += f"- File size: {len(response.content)} bytes\n"
                debug_info += f"- Signature length: {len(signature)} bytes\n"
                debug_info += f"- Public key length: {len(public_key_buyer)} bytes\n"
                debug_info += f"- Data to verify length: {len(docx_bytes_to_verify)} bytes\n"
                debug_info += f"- Data hash: {hashlib.sha256(docx_bytes_to_verify).hexdigest()[:16]}"
                
                print(debug_info)
                QMessageBox.warning(self, "Lỗi", f"Xác thực thất bại: Chữ ký của người mua không hợp lệ.\n\n{debug_info}")

        except Exception as e:
            print(f"DEBUG: Exception trong download_and_verify_order: {str(e)}")
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi tải và xác thực order: {str(e)}")

    def sign_downloaded_order_as_invoice(self, order_docx_path, buyer_name):
        """Người bán ký Order đã tải về và tạo Invoice với mã QR."""
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu của bạn:", QLineEdit.Password
        )
        if not ok or not password:
            return
        
        try:
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key của bạn.")
                return

            # Đọc order để lấy metadata
            doc_order = Document(order_docx_path)
            metadata_order = self._extract_metadata_from_docx(doc_order)
            if not metadata_order:
                QMessageBox.warning(self, "Lỗi", "Không thể đọc metadata từ order")
                return
            
            # Tạo hóa đơn mới từ dữ liệu order
            doc_invoice = Document()
            doc_invoice.add_heading('HÓA ĐƠN', 0)
            
            # Thông tin hóa đơn
            doc_invoice.add_heading('Thông tin Hóa đơn:', level=1)
            invoice_info_table = doc_invoice.add_table(rows=5, cols=2)
            invoice_info_table.style = 'Table Grid'
            
            invoice_number = f"INV{datetime.now().strftime('%Y%m%d%H%M%S')}"
            invoice_date = datetime.now().strftime('%d/%m/%Y')
            due_date = (datetime.now() + timedelta(days=30)).strftime('%d/%m/%Y')
            
            invoice_info_table.cell(0, 0).text = "Số hóa đơn (Invoice Number):"
            invoice_info_table.cell(0, 1).text = invoice_number
            invoice_info_table.cell(1, 0).text = "Ngày hóa đơn (Invoice Date):"
            invoice_info_table.cell(1, 1).text = invoice_date
            invoice_info_table.cell(2, 0).text = "Ngày đáo hạn (Due Date):"
            invoice_info_table.cell(2, 1).text = due_date
            invoice_info_table.cell(3, 0).text = "Điều khoản thanh toán:"
            invoice_info_table.cell(3, 1).text = "Net 30"
            invoice_info_table.cell(4, 0).text = "Tình trạng thanh toán:"
            invoice_info_table.cell(4, 1).text = "Unpaid"
            
            # Thông tin bên bán
            doc_invoice.add_heading('Thông tin Bên bán (Nhà cung cấp):', level=1)
            seller_info_table = doc_invoice.add_table(rows=7, cols=2)
            seller_info_table.style = 'Table Grid'
            
            seller_info_table.cell(0, 0).text = "Tên công ty/Cửa hàng:"
            seller_info_table.cell(0, 1).text = f"Công ty {self.current_user}"
            seller_info_table.cell(1, 0).text = "Địa chỉ:"
            seller_info_table.cell(1, 1).text = "456 Đường XYZ, Quận 3, TP.HCM"
            seller_info_table.cell(2, 0).text = "Số điện thoại:"
            seller_info_table.cell(2, 1).text = "0281234567"
            seller_info_table.cell(3, 0).text = "Email:"
            seller_info_table.cell(3, 1).text = f"{self.current_user}@company.com"
            seller_info_table.cell(4, 0).text = "Mã số thuế:"
            seller_info_table.cell(4, 1).text = "9876543210"
            seller_info_table.cell(5, 0).text = "Tài khoản ngân hàng:"
            seller_info_table.cell(5, 1).text = "1234567890"
            seller_info_table.cell(6, 0).text = "Ngân hàng:"
            seller_info_table.cell(6, 1).text = "Vietcombank - Chi nhánh TP.HCM"
            
            # Thông tin khách hàng
            doc_invoice.add_heading('Thông tin Khách hàng (Bên mua):', level=1)
            customer_info_table = doc_invoice.add_table(rows=5, cols=2)
            customer_info_table.style = 'Table Grid'
            
            customer_info = metadata_order.get('customer_info', {})
            
            customer_info_table.cell(0, 0).text = "Tên khách hàng/Tên công ty:"
            customer_info_table.cell(0, 1).text = customer_info.get('name', buyer_name)
            customer_info_table.cell(1, 0).text = "Địa chỉ thanh toán:"
            customer_info_table.cell(1, 1).text = customer_info.get('address', "123 Đường ABC, Quận 1, TP.HCM")
            customer_info_table.cell(2, 0).text = "Số điện thoại:"
            customer_info_table.cell(2, 1).text = customer_info.get('phone', "0901234567")
            customer_info_table.cell(3, 0).text = "Email:"
            customer_info_table.cell(3, 1).text = customer_info.get('email', f"{buyer_name}@email.com")
            customer_info_table.cell(4, 0).text = "Mã số thuế:"
            customer_info_table.cell(4, 1).text = customer_info.get('tax_code', "0123456789")
            
            # Danh sách sản phẩm từ order
            doc_invoice.add_heading('Thông tin Sản phẩm/Dịch vụ:', level=1)
            table = doc_invoice.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'STT'
            header_cells[1].text = 'Tên sản phẩm/Dịch vụ'
            header_cells[2].text = 'Mã sản phẩm'
            header_cells[3].text = 'Số lượng'
            header_cells[4].text = 'Đơn giá (VNĐ)'
            header_cells[5].text = 'Thành tiền (VNĐ)'
            
            # Thêm dữ liệu sản phẩm từ order metadata
            order_items = metadata_order.get('order_items', [])
            subtotal = 0
            
            for i, item in enumerate(order_items):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = item['product_name']
                row_cells[2].text = f"SP{item['product_id']:04d}"
                row_cells[3].text = str(item['quantity'])
                row_cells[4].text = f"{item['price']:,}"
                row_cells[5].text = f"{item['total']:,}"
                subtotal += item['total']
            
            # Tổng kết hóa đơn
            doc_invoice.add_heading('Tổng kết Hóa đơn:', level=1)
            summary_table = doc_invoice.add_table(rows=8, cols=2)
            summary_table.style = 'Table Grid'
            
            # Lấy thông tin từ order hoặc sử dụng giá trị mặc định
            shipping_fee = metadata_order.get('shipping_fee', 50000)
            discount = metadata_order.get('discount', 0)
            tax_amount = metadata_order.get('tax_amount', subtotal * 0.1)
            total_amount = metadata_order.get('total_amount', subtotal + shipping_fee + tax_amount - discount)
            amount_paid = 0  # Chưa thanh toán
            balance_due = total_amount - amount_paid
            
            summary_table.cell(0, 0).text = "Tổng tiền hàng:"
            summary_table.cell(0, 1).text = f"{subtotal:,} VNĐ"
            summary_table.cell(1, 0).text = "Chiết khấu:"
            summary_table.cell(1, 1).text = f"{discount:,} VNĐ"
            summary_table.cell(2, 0).text = "Phí vận chuyển:"
            summary_table.cell(2, 1).text = f"{shipping_fee:,} VNĐ"
            summary_table.cell(3, 0).text = "Thuế VAT (10%):"
            summary_table.cell(3, 1).text = f"{tax_amount:,} VNĐ"
            summary_table.cell(4, 0).text = "Tổng cộng phải trả:"
            summary_table.cell(4, 1).text = f"{total_amount:,} VNĐ"
            summary_table.cell(5, 0).text = "Số tiền đã nhận/Đã thanh toán:"
            summary_table.cell(5, 1).text = f"{amount_paid:,} VNĐ"
            summary_table.cell(6, 0).text = "Số tiền còn lại phải trả:"
            summary_table.cell(6, 1).text = f"{balance_due:,} VNĐ"
            summary_table.cell(7, 0).text = "Phương thức thanh toán:"
            summary_table.cell(7, 1).text = "Chuyển khoản ngân hàng"
            
            # Ghi chú/Điều khoản
            doc_invoice.add_heading('Ghi chú/Điều khoản:', level=1)
            doc_invoice.add_paragraph("• Thanh toán trong vòng 30 ngày kể từ ngày hóa đơn")
            doc_invoice.add_paragraph("• Bảo hành sản phẩm theo chính sách của nhà sản xuất")
            doc_invoice.add_paragraph("• Cảm ơn quý khách đã tin tưởng và sử dụng dịch vụ của chúng tôi")
            doc_invoice.add_paragraph("• Mọi thắc mắc vui lòng liên hệ: 0281234567")
            
            # Thêm QR code của người bán
            timestamp_seller = datetime.now().isoformat()
            qr_image = self.create_qr_code(self.current_user, timestamp_seller)
            self.add_qr_to_docx(doc_invoice, qr_image)
            
            # Thêm metadata cho hóa đơn
            invoice_metadata = {
                'signer': self.current_user,
                'sign_time': timestamp_seller,
                'invoice_number': invoice_number,
                'invoice_date': invoice_date,
                'due_date': due_date,
                'payment_terms': 'Net 30',
                'payment_status': 'Unpaid',
                'buyer_name': buyer_name,
                'order_items': order_items,
                'subtotal': subtotal,
                'shipping_fee': shipping_fee,
                'tax_amount': tax_amount,
                'discount': discount,
                'total_amount': total_amount,
                'amount_paid': amount_paid,
                'balance_due': balance_due,
                'original_order_signature': metadata_order.get('signature'),
                'seller_info': {
                    'name': f"Công ty {self.current_user}",
                    'address': "456 Đường XYZ, Quận 3, TP.HCM",
                    'phone': "0281234567",
                    'email': f"{self.current_user}@company.com",
                    'tax_code': "9876543210",
                    'bank_account': "1234567890",
                    'bank_name': "Vietcombank - Chi nhánh TP.HCM"
                },
                'customer_info': customer_info
            }
            self._store_metadata_in_docx(doc_invoice, invoice_metadata)
            
            # Lấy bytes để ký
            buffer_to_sign = io.BytesIO()
            doc_invoice.save(buffer_to_sign)
            bytes_to_sign = buffer_to_sign.getvalue()
            
            # Người bán ký hóa đơn
            signature_seller = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64_seller = base64.b64encode(signature_seller).decode('utf-8')
            
            # Cập nhật metadata với chữ ký
            invoice_metadata['signature'] = signature_b64_seller
            invoice_metadata['signed_data_bytes'] = base64.b64encode(bytes_to_sign).decode('utf-8')
            invoice_metadata['signed_data_hash'] = hashlib.sha256(bytes_to_sign).hexdigest()
            self._store_metadata_in_docx(doc_invoice, invoice_metadata)
            
            # Lưu file hóa đơn
            final_invoice_docx_path = os.path.join(os.path.expanduser("~"), "Downloads", f"invoice_{buyer_name}_from_{self.current_user}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc_invoice.save(final_invoice_docx_path)
            
            # Upload hóa đơn lên server
            with open(final_invoice_docx_path, 'rb') as f:
                files = {'pdf': f}
                data = {
                    'signer_name': self.current_user,
                    'signature': signature_b64_seller,
                    'timestamp': timestamp_seller
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-invoice/",
                    files=files,
                    data=data
                )
                response.raise_for_status()

            QMessageBox.information(
                self, "Thành công",
                f"Đã tạo hóa đơn thành công từ order!\nSố hóa đơn: {invoice_number}\nHóa đơn được lưu tại:\n{final_invoice_docx_path}"
            )
            self.refresh_order_list()
            self.refresh_my_invoice_list()
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi tạo hóa đơn: {str(e)}")

    def show_register_dialog(self):
        """Hiển thị dialog đăng ký"""
        dialog = RegisterDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            self.register_user(data)

    def register_user(self, data):
        """Xử lý đăng ký người dùng"""
        username = data['username']
        password = data['password']
        confirm_password = data['confirm_password']
        role = data['role']
        
        if not username or not password or not role:
            QMessageBox.warning(self, "Lỗi", "Vui lòng điền đầy đủ thông tin và chọn vai trò")
            return
        
        if password != confirm_password:
            QMessageBox.warning(self, "Lỗi", "Mật khẩu xác nhận không khớp")
            return
        
        try:
            # Tạo cặp khóa mới
            public_key, private_key = ML_DSA_44.keygen()
            
            # Lưu khóa
            if not self.save_keys(public_key, private_key, password, username):
                QMessageBox.critical(self, "Lỗi", "Không thể lưu khóa")
                return
            
            # Tạo payload cho proof of possession
            payload = {
                'username': username,
                'role': role,
                'timestamp': datetime.now().isoformat()
            }
            
            # Ký payload để chứng minh sở hữu private key
            signature = ML_DSA_44.sign(private_key, json.dumps(payload, sort_keys=True).encode())
            
            # Gửi request đăng ký
            response = requests.post(
                f"{self.API_BASE_URL}/register/",
                json={
                    'public_key': base64.b64encode(public_key).decode('utf-8'),
                    'payload': payload,
                    'signature': base64.b64encode(signature).decode('utf-8')
                }
            )
            
            if response.status_code == 200:
                response_data = response.json()
                
                # Lưu CA info
                ca_info = {
                    'ca_public_key': response_data['ca_public_key'],
                    'ca_public_key_hash': response_data['ca_public_key_hash']
                }
                
                ca_info_path = os.path.join(self.keys_dir, 'ca_info.json')
                with open(ca_info_path, 'w') as f:
                    json.dump(ca_info, f)
                
                QMessageBox.information(
                    self, "Thành công",
                    f"Đăng ký thành công!\nUsername: {username}\nVai trò: {role}\n\nVui lòng đăng nhập để sử dụng hệ thống."
                )
            else:
                error_msg = response.json().get('error', 'Lỗi không xác định')
                QMessageBox.critical(self, "Lỗi", f"Đăng ký thất bại: {error_msg}")
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi đăng ký: {str(e)}")

    # Product management methods
    def add_product(self):
        """Thêm sản phẩm mới"""
        if not self.current_user:
            QMessageBox.warning(self, "Lỗi", "Vui lòng đăng nhập trước")
            return
            
        product_name = self.product_name_input.text().strip()
        price_text = self.product_price_input.text().strip()
        
        if not product_name or not price_text:
            QMessageBox.warning(self, "Lỗi", "Vui lòng nhập đầy đủ tên sản phẩm và giá")
            return
            
        try:
            price = float(price_text)
            if price <= 0:
                QMessageBox.warning(self, "Lỗi", "Giá phải lớn hơn 0")
                return
        except ValueError:
            QMessageBox.warning(self, "Lỗi", "Giá không hợp lệ")
            return
            
        try:
            response = requests.post(
                f"{self.API_BASE_URL}/add-product/",
                data={
                    'seller_name': self.current_user,
                    'product_name': product_name,
                    'price': price
                }
            )
            
            if response.status_code == 200:
                QMessageBox.information(self, "Thành công", "Đã thêm sản phẩm thành công!")
                self.product_name_input.clear()
                self.product_price_input.clear()
                self.refresh_my_products()
            else:
                error_msg = response.json().get('error', 'Lỗi không xác định')
                QMessageBox.critical(self, "Lỗi", f"Thêm sản phẩm thất bại: {error_msg}")
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi thêm sản phẩm: {str(e)}")

    def refresh_my_products(self):
        """Làm mới danh sách sản phẩm của tôi"""
        if not self.current_user:
            return
            
        try:
            response = requests.get(
                f"{self.API_BASE_URL}/list-my-products/",
                params={'seller_name': self.current_user}
            )
            response.raise_for_status()
            
            products = response.json()
            self.my_products_list.clear()
            
            for product in products:
                item = QListWidgetItem(
                    f"{product['product_name']} - {product['price']:,} VNĐ"
                )
                item.setData(Qt.UserRole, product)
                self.my_products_list.addItem(item)
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách sản phẩm: {str(e)}")

    def refresh_products(self):
        """Làm mới danh sách tất cả sản phẩm"""
        try:
            response = requests.get(f"{self.API_BASE_URL}/list-products/")
            response.raise_for_status()
            
            products = response.json()
            self.products_list.clear()
            
            for product in products:
                item = QListWidgetItem(
                    f"{product['product_name']} - {product['price']:,} VNĐ (Bán bởi: {product['seller_name']})"
                )
                item.setData(Qt.UserRole, product)
                self.products_list.addItem(item)
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách sản phẩm: {str(e)}")

    def select_product_for_order(self, item):
        """Chọn sản phẩm để thêm vào giỏ hàng"""
        product = item.data(Qt.UserRole)
        
        quantity, ok = QInputDialog.getInt(
            self, "Chọn số lượng", 
            f"Số lượng {product['product_name']}:", 
            value=1, min=1, max=100
        )
        
        if ok:
            # Thêm vào giỏ hàng
            cart_item = {
                'product': product,
                'quantity': quantity,
                'total': product['price'] * quantity
            }
            
            # Kiểm tra xem sản phẩm đã có trong giỏ hàng chưa
            for i in range(self.cart_list.count()):
                existing_item = self.cart_list.item(i)
                existing_data = existing_item.data(Qt.UserRole)
                if existing_data['product']['id'] == product['id']:
                    # Cập nhật số lượng
                    existing_data['quantity'] += quantity
                    existing_data['total'] = product['price'] * existing_data['quantity']
                    existing_item.setText(
                        f"{product['product_name']} x{existing_data['quantity']} - {existing_data['total']:,}"
                    )
                    existing_item.setData(Qt.UserRole, existing_data)
                    self.update_cart_total()
                    return
            
            # Thêm mới vào giỏ hàng
            item = QListWidgetItem(
                f"{product['product_name']} x{quantity} - {cart_item['total']:,}"
            )
            item.setData(Qt.UserRole, cart_item)
            self.cart_list.addItem(item)
            self.update_cart_total()

    def update_cart_total(self):
        """Cập nhật tổng tiền giỏ hàng"""
        total = 0
        for i in range(self.cart_list.count()):
            item = self.cart_list.item(i)
            cart_data = item.data(Qt.UserRole)
            total += cart_data['total']
        
        self.total_price_label.setText(f"Tổng tiền: {total:,} VNĐ")

    def clear_cart(self):
        """Xóa giỏ hàng"""
        self.cart_list.clear()
        self.update_cart_total()

    def create_order_from_cart(self):
        """Tạo đơn hàng từ giỏ hàng"""
        if self.cart_list.count() == 0:
            QMessageBox.warning(self, "Lỗi", "Giỏ hàng trống")
            return
            
        if not self.current_user:
            QMessageBox.warning(self, "Lỗi", "Vui lòng đăng nhập trước")
            return
            
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu:", QLineEdit.Password
        )
        if not ok or not password:
            return
            
        try:
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key")
                return
            
            # Tạo file DOCX
            doc = Document()
            doc.add_heading('ĐƠN HÀNG', 0)
            
            # Thông tin đơn hàng
            doc.add_heading('Thông tin Đơn hàng:', level=1)
            order_info_table = doc.add_table(rows=4, cols=2)
            order_info_table.style = 'Table Grid'
            
            order_number = f"ORD{datetime.now().strftime('%Y%m%d%H%M%S')}"
            order_date = datetime.now().strftime('%d/%m/%Y')
            
            order_info_table.cell(0, 0).text = "Mã số Đơn hàng:"
            order_info_table.cell(0, 1).text = order_number
            order_info_table.cell(1, 0).text = "Ngày đặt hàng:"
            order_info_table.cell(1, 1).text = order_date
            order_info_table.cell(2, 0).text = "Kênh đặt hàng:"
            order_info_table.cell(2, 1).text = "Hệ thống điện tử"
            order_info_table.cell(3, 0).text = "Phương thức thanh toán:"
            order_info_table.cell(3, 1).text = "Chuyển khoản ngân hàng"
            
            # Thông tin khách hàng
            doc.add_heading('Thông tin Khách hàng (Bên mua):', level=1)
            customer_info_table = doc.add_table(rows=5, cols=2)
            customer_info_table.style = 'Table Grid'
            
            customer_info_table.cell(0, 0).text = "Tên khách hàng:"
            customer_info_table.cell(0, 1).text = self.current_user
            customer_info_table.cell(1, 0).text = "Địa chỉ giao hàng:"
            customer_info_table.cell(1, 1).text = "123 Đường ABC, Quận 1, TP.HCM"
            customer_info_table.cell(2, 0).text = "Số điện thoại:"
            customer_info_table.cell(2, 1).text = "0901234567"
            customer_info_table.cell(3, 0).text = "Email:"
            customer_info_table.cell(3, 1).text = f"{self.current_user}@email.com"
            customer_info_table.cell(4, 0).text = "Mã số thuế:"
            customer_info_table.cell(4, 1).text = "0123456789"
            
            # Danh sách sản phẩm
            doc.add_heading('Thông tin Sản phẩm/Dịch vụ:', level=1)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'STT'
            header_cells[1].text = 'Tên sản phẩm/Dịch vụ'
            header_cells[2].text = 'Số lượng'
            header_cells[3].text = 'Đơn giá (VNĐ)'
            header_cells[4].text = 'Thành tiền (VNĐ)'
            
            # Thêm dữ liệu sản phẩm
            subtotal = 0
            for i in range(self.cart_list.count()):
                item = self.cart_list.item(i)
                cart_data = item.data(Qt.UserRole)
                product = cart_data['product']
                
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = product['product_name']
                row_cells[2].text = str(cart_data['quantity'])
                row_cells[3].text = f"{product['price']:,}"
                row_cells[4].text = f"{cart_data['total']:,}"
                
                subtotal += cart_data['total']
            
            # Tổng kết đơn hàng
            doc.add_heading('Tổng kết Đơn hàng:', level=1)
            summary_table = doc.add_table(rows=6, cols=2)
            summary_table.style = 'Table Grid'
            
            shipping_fee = 50000  # Phí vận chuyển
            discount = 0  # Chiết khấu
            tax_rate = 0.1  # Thuế VAT 10%
            tax_amount = subtotal * tax_rate
            total_amount = subtotal + shipping_fee + tax_amount - discount
            
            summary_table.cell(0, 0).text = "Tổng tiền hàng (Subtotal):"
            summary_table.cell(0, 1).text = f"{subtotal:,} VNĐ"
            summary_table.cell(1, 0).text = "Chiết khấu/Mã giảm giá:"
            summary_table.cell(1, 1).text = f"{discount:,} VNĐ"
            summary_table.cell(2, 0).text = "Phí vận chuyển:"
            summary_table.cell(2, 1).text = f"{shipping_fee:,} VNĐ"
            summary_table.cell(3, 0).text = "Thuế VAT (10%):"
            summary_table.cell(3, 1).text = f"{tax_amount:,} VNĐ"
            summary_table.cell(4, 0).text = "Tổng cộng phải trả:"
            summary_table.cell(4, 1).text = f"{total_amount:,} VNĐ"
            summary_table.cell(5, 0).text = "Phương thức thanh toán:"
            summary_table.cell(5, 1).text = "Chuyển khoản ngân hàng"
            
            # Thêm QR code
            qr_image = self.create_qr_code(self.current_user, datetime.now().isoformat())
            self.add_qr_to_docx(doc, qr_image)
            
            # Thêm metadata
            metadata = {
                'buyer': self.current_user,
                'sign_time': datetime.now().isoformat(),
                'order_number': order_number,
                'order_date': order_date,
                'order_items': [
                    {
                        'product_id': item.data(Qt.UserRole)['product']['id'],
                        'product_name': item.data(Qt.UserRole)['product']['product_name'],
                        'quantity': item.data(Qt.UserRole)['quantity'],
                        'price': item.data(Qt.UserRole)['product']['price'],
                        'total': item.data(Qt.UserRole)['total']
                    }
                    for i in range(self.cart_list.count())
                    for item in [self.cart_list.item(i)]
                ],
                'subtotal': subtotal,
                'shipping_fee': shipping_fee,
                'tax_amount': tax_amount,
                'discount': discount,
                'total_amount': total_amount,
                'customer_info': {
                    'name': self.current_user,
                    'address': "123 Đường ABC, Quận 1, TP.HCM",
                    'phone': "0901234567",
                    'email': f"{self.current_user}@email.com",
                    'tax_code': "0123456789"
                }
            }
            self._store_metadata_in_docx(doc, metadata)
            
            # Lấy bytes để ký
            buffer_to_sign = io.BytesIO()
            doc.save(buffer_to_sign)
            bytes_to_sign = buffer_to_sign.getvalue()
            
            # Ký
            signature = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64 = base64.b64encode(signature).decode('utf-8')
            
            # Cập nhật metadata với chữ ký
            metadata['signature'] = signature_b64
            metadata['signed_data_bytes'] = base64.b64encode(bytes_to_sign).decode('utf-8')
            metadata['signed_data_hash'] = hashlib.sha256(bytes_to_sign).hexdigest()
            self._store_metadata_in_docx(doc, metadata)
            
            # Lưu file
            order_path = os.path.join(os.path.expanduser("~"), "Downloads", f"order_{self.current_user}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(order_path)
            
            # Upload lên server
            with open(order_path, 'rb') as f:
                files = {'pdf': f}
                data = {
                    'buyer_name': self.current_user,
                    'signature': signature_b64,
                    'timestamp': datetime.now().isoformat()
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-order/",
                    files=files,
                    data=data
                )
                response.raise_for_status()
            
            QMessageBox.information(
                self, "Thành công",
                f"Đã tạo đơn hàng thành công!\nMã đơn hàng: {order_number}\nFile được lưu tại: {order_path}"
            )
            
            # Xóa giỏ hàng
            self.clear_cart()
            self.refresh_my_order_list()
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi tạo đơn hàng: {str(e)}")

    def refresh_my_invoice_list(self):
        """Làm mới danh sách hóa đơn của tôi"""
        if not self.current_user:
            return
            
        try:
            response = requests.get(
                f"{self.API_BASE_URL}/list-my-invoices/",
                params={'seller_name': self.current_user}
            )
            response.raise_for_status()
            
            invoices = response.json()
            self.my_invoice_list.clear()
            
            for invoice in invoices:
                item = QListWidgetItem(
                    f"Hóa đơn {invoice['id']} - {invoice['timestamp']}"
                )
                item.setData(Qt.UserRole, invoice)
                self.my_invoice_list.addItem(item)
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách hóa đơn: {str(e)}")

    def refresh_my_order_list(self):
        """Làm mới danh sách order của tôi"""
        if not self.current_user:
            return
            
        try:
            response = requests.get(
                f"{self.API_BASE_URL}/list-my-orders/",
                params={'buyer_name': self.current_user}
            )
            response.raise_for_status()
            
            orders = response.json()
            self.my_order_list.clear()
            
            for order in orders:
                item = QListWidgetItem(
                    f"Order {order['id']} - {order['timestamp']}"
                )
                item.setData(Qt.UserRole, order)
                self.my_order_list.addItem(item)
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách order: {str(e)}")


class RegisterDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        """Khởi tạo giao diện dialog đăng ký"""
        self.setWindowTitle('Đăng ký người dùng mới')
        self.setModal(True)
        self.setFixedSize(300, 200)

        layout = QVBoxLayout(self)

        # Form layout
        form_layout = QFormLayout()

        self.username_input = QLineEdit()
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        self.confirm_password_input = QLineEdit()
        self.confirm_password_input.setEchoMode(QLineEdit.Password)

        form_layout.addRow("Username:", self.username_input)
        form_layout.addRow("Mật khẩu:", self.password_input)
        form_layout.addRow("Xác nhận mật khẩu:", self.confirm_password_input)

        # Role selection
        self.role_combo = QComboBox()
        self.role_combo.addItems(['buyer', 'seller'])
        form_layout.addRow("Vai trò:", self.role_combo)

        layout.addLayout(form_layout)

        # Buttons
        button_layout = QHBoxLayout()
        register_btn = QPushButton("Đăng ký")
        register_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton("Hủy")
        cancel_btn.clicked.connect(self.reject)

        button_layout.addWidget(register_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)

    def get_data(self):
        """Lấy dữ liệu từ dialog"""
        return {
            'username': self.username_input.text().strip(),
            'password': self.password_input.text(),
            'confirm_password': self.confirm_password_input.text(),
            'role': self.role_combo.currentText()
        }


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = PDFVerifier()
    window.show()
    sys.exit(app.exec_())