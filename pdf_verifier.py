import sys
import os
import json
import logging
import requests
import base64
import io
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QLineEdit, QTabWidget,
    QMessageBox, QTextEdit, QGroupBox, QFormLayout, QInputDialog,
    QListWidget, QListWidgetItem, QDialog, QComboBox
)
from PyQt5.QtCore import Qt
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from Crypto.Protocol.KDF import PBKDF2
from dilithium_py.ml_dsa import ML_DSA_44
from PyPDF2 import PdfReader, PdfWriter
from PyPDF2.generic import NameObject, TextStringObject
from datetime import datetime
import qrcode
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import tempfile
import hashlib
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cấu hình logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def verify_certificate(certificate):
    """Verify a certificate using CA's public key"""
    try:
        # Read CA info from file
        ca_info_path = os.path.join(os.path.expanduser("~"), ".pdf_verifier", "keys", "ca_info.json")
        if not os.path.exists(ca_info_path):
            print("CA info file not found")
            return False
            
        with open(ca_info_path, 'r') as f:
            ca_info = json.load(f)
        
        # Use the verified CA public key
        ca_public_key = base64.b64decode(ca_info['ca_public_key'])
        
        # Extract payload and signature from certificate
        payload = certificate['payload']
        signature = base64.b64decode(certificate['signature'])
        
        # Verify signature
        return ML_DSA_44.verify(ca_public_key, json.dumps(payload, sort_keys=True).encode(), signature)
    except Exception as e:
        print(f"Error verifying certificate: {e}")
        return False

class PDFVerifier(QMainWindow):
    def __init__(self):
        super().__init__()
        self.API_BASE_URL = "http://localhost:8000"
        self.current_user = None
        self.current_user_role = None
        self.keys_dir = os.path.join(os.path.expanduser("~"), ".pdf_verifier", "keys")
        os.makedirs(self.keys_dir, exist_ok=True)
        self.init_ui()

    def init_ui(self):
        """Khởi tạo giao diện"""
        self.setWindowTitle('Hệ thống Ký và Xác thực PDF')
        self.setGeometry(100, 100, 1000, 600)

        # Widget chính
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        self.main_layout = QVBoxLayout(main_widget)

        # Tạo login widget
        self.login_widget = QWidget()
        self.init_login_ui()
        self.main_layout.addWidget(self.login_widget)

        # Tạo main content widget (sẽ ẩn ban đầu)
        self.content_widget = QWidget()
        self.init_content_ui()
        self.content_widget.hide()
        self.main_layout.addWidget(self.content_widget)

    def init_login_ui(self):
        """Khởi tạo giao diện đăng nhập"""
        login_layout = QVBoxLayout(self.login_widget)
        
        # Tạo group box cho phần đăng nhập
        login_group = QGroupBox("Đăng nhập")
        login_form = QFormLayout()
        
        self.username_input = QLineEdit()
        login_form.addRow("Username:", self.username_input)
        
        # Thêm layout ngang cho các nút
        button_layout = QHBoxLayout()
        
        login_btn = QPushButton("Đăng nhập")
        login_btn.clicked.connect(self.login)
        button_layout.addWidget(login_btn)
        
        register_btn = QPushButton("Đăng ký")
        register_btn.clicked.connect(self.show_register_dialog)
        button_layout.addWidget(register_btn)
        
        login_form.addRow("", button_layout)
        login_group.setLayout(login_form)
        
        # Thêm vào layout chính
        login_layout.addWidget(login_group)
        login_layout.addStretch()

    def init_content_ui(self):
        """Khởi tạo giao diện chính sau khi đăng nhập"""
        content_layout = QVBoxLayout(self.content_widget)
        
        # Thêm header với thông tin người dùng và nút đăng xuất
        header_layout = QHBoxLayout()
        self.user_label = QLabel()
        header_layout.addWidget(self.user_label)
        
        logout_btn = QPushButton("Đăng xuất")
        logout_btn.clicked.connect(self.logout)
        header_layout.addWidget(logout_btn)
        
        content_layout.addLayout(header_layout)
        
        # Tab widget
        self.tabs = QTabWidget()
        
        # Tab Bên bán (khởi tạo nhưng chưa thêm vào tabs)
        self.seller_tab = QWidget()
        seller_layout = QVBoxLayout()
        
        # Sign section
        sign_group = QGroupBox("Ký và Upload PDF")
        sign_layout = QVBoxLayout()
        
        self.sign_path = QLineEdit()
        self.sign_path.setReadOnly(True)
        sign_select_btn = QPushButton("Chọn file DOCX")
        sign_select_btn.clicked.connect(self.select_sign_file)
        
        sign_layout.addWidget(QLabel("File DOCX:"))
        sign_layout.addWidget(self.sign_path)
        sign_layout.addWidget(sign_select_btn)
        
        sign_btn = QPushButton("Ký và Upload DOCX")
        sign_btn.clicked.connect(self.sign_and_upload_docx)
        sign_layout.addWidget(sign_btn)
        
        sign_group.setLayout(sign_layout)
        seller_layout.addWidget(sign_group)
        
        # Thêm danh sách order từ bên mua
        order_group = QGroupBox("Danh sách order từ bên mua")
        order_layout = QVBoxLayout()
        self.order_list = QListWidget()
        self.order_list.itemDoubleClicked.connect(self.download_and_verify_order)
        order_layout.addWidget(self.order_list)
        refresh_order_btn = QPushButton("Làm mới danh sách order")
        refresh_order_btn.clicked.connect(self.refresh_order_list)
        order_layout.addWidget(refresh_order_btn)
        order_group.setLayout(order_layout)
        seller_layout.addWidget(order_group)
        
        self.seller_tab.setLayout(seller_layout)

        # Tab Bên mua (khởi tạo nhưng chưa thêm vào tabs)
        self.buyer_tab = QWidget()
        buyer_layout = QVBoxLayout()
        
        # Invoice list section
        invoice_group = QGroupBox("Danh sách hóa đơn")
        invoice_layout = QVBoxLayout()
        
        self.invoice_list = QListWidget()
        self.invoice_list.itemDoubleClicked.connect(self.download_and_verify_invoice)
        invoice_layout.addWidget(self.invoice_list)
        
        refresh_btn = QPushButton("Làm mới danh sách")
        refresh_btn.clicked.connect(self.refresh_invoice_list)
        invoice_layout.addWidget(refresh_btn)
        
        invoice_group.setLayout(invoice_layout)
        buyer_layout.addWidget(invoice_group)
        
        # Thêm chức năng upload order PDF
        order_upload_group = QGroupBox("Ký và Upload Order PDF")
        order_upload_layout = QVBoxLayout()
        self.order_sign_path = QLineEdit()
        self.order_sign_path.setReadOnly(True)
        order_sign_select_btn = QPushButton("Chọn file Order DOCX")
        order_sign_select_btn.clicked.connect(self.select_order_sign_file)
        order_upload_layout.addWidget(QLabel("File Order DOCX:"))
        order_upload_layout.addWidget(self.order_sign_path)
        order_upload_layout.addWidget(order_sign_select_btn)
        order_sign_btn = QPushButton("Ký và Upload Order DOCX")
        order_sign_btn.clicked.connect(self.sign_and_upload_order_docx)
        order_upload_layout.addWidget(order_sign_btn)
        order_upload_group.setLayout(order_upload_layout)
        buyer_layout.addWidget(order_upload_group)
        
        # Thêm danh sách order đã gửi
        my_order_group = QGroupBox("Order đã gửi")
        my_order_layout = QVBoxLayout()
        self.my_order_list = QListWidget()
        my_order_layout.addWidget(self.my_order_list)
        my_order_group.setLayout(my_order_layout)
        buyer_layout.addWidget(my_order_group)
        
        self.buyer_tab.setLayout(buyer_layout)

        content_layout.addWidget(self.tabs)

    def login(self):
        """Đăng nhập với username"""
        username = self.username_input.text().strip()
        if not username:
            QMessageBox.warning(self, "Lỗi", "Vui lòng nhập username")
            return
            
        try:
            # Kiểm tra user tồn tại và lấy vai trò
            response = requests.get(
                f"{self.API_BASE_URL}/check-user",
                params={"username": username}
            )
            response_data = response.json()
            
            if response.status_code == 200 and response_data.get('message') == 'User exists':
                self.current_user = username
                self.current_user_role = response_data.get('role')
                self.user_label.setText(f"Đã đăng nhập: {username} ({self.current_user_role})")
                
                # Ẩn login widget và hiện content widget
                self.login_widget.hide()
                self.content_widget.show()
                
                self.update_ui_for_role()
                
                QMessageBox.information(self, "Thành công", f"Đăng nhập thành công: {username}")
            else:
                QMessageBox.warning(self, "Lỗi", response_data.get('message', 'Username không tồn tại hoặc lỗi server'))
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi đăng nhập: {str(e)}")

    def logout(self):
        """Đăng xuất"""
        self.current_user = None
        self.current_user_role = None
        self.username_input.clear()
        
        # Xóa tất cả các tab khi đăng xuất
        self.tabs.clear()

        # Ẩn content widget và hiện login widget
        self.content_widget.hide()
        self.login_widget.show()
        
        QMessageBox.information(self, "Thông báo", "Đã đăng xuất thành công")

    def update_ui_for_role(self):
        """Cập nhật giao diện dựa trên vai trò của người dùng"""
        # Xóa tất cả các tab hiện có
        self.tabs.clear() 

        if self.current_user_role == 'seller':
            self.tabs.addTab(self.seller_tab, "Bên bán")
            self.tabs.setCurrentWidget(self.seller_tab)
            self.refresh_order_list()
        elif self.current_user_role == 'buyer':
            self.tabs.addTab(self.buyer_tab, "Bên mua")
            self.tabs.setCurrentWidget(self.buyer_tab)
            self.refresh_invoice_list()
            self.refresh_my_order_list()
        else:
            QMessageBox.warning(self, "Lỗi", "Vai trò người dùng không xác định. Vui lòng liên hệ quản trị viên.")

    def select_sign_file(self):
        """Chọn file DOCX để ký"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Chọn file DOCX", "", "Word Documents (*.docx)"
        )
        if file_path:
            self.sign_path.setText(file_path)

    def sign_and_upload_docx(self):
        """Ký và upload file DOCX"""
        if not self.current_user:
            QMessageBox.warning(self, "Lỗi", "Vui lòng đăng nhập trước")
            return
        file_path = self.sign_path.text()
        if not file_path:
            QMessageBox.warning(self, "Lỗi", "Vui lòng chọn file DOCX")
            return
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu:", QLineEdit.Password
        )
        if not ok or not password:
            return
        try:
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key")
                return
            
            # Lấy public key từ server
            try:
                response = requests.get(
                    f"{self.API_BASE_URL}/get-public-key/",
                    params={'signer_name': self.current_user}
                )
                response.raise_for_status()
                public_key = base64.b64decode(response.json()['public_key'])
            except Exception as e:
                QMessageBox.critical(self, "Lỗi", f"Không thể lấy public key từ server: {str(e)}")
                return

            # Bước 1: Đọc DOCX gốc và tạo metadata
            doc = Document(file_path)
            
            # Tạo metadata cho document
            metadata = {
                'signer': self.current_user,
                'signature': '__SIGNATURE_PLACEHOLDER__',  # Placeholder
                'sign_time': datetime.now().isoformat()
            }
            
            # Thêm metadata vào document properties
            doc.core_properties.author = self.current_user
            doc.core_properties.comments = json.dumps(metadata)

            # Lưu DOCX với metadata vào file tạm thời
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx_with_metadata:
                doc.save(tmp_docx_with_metadata.name)
                tmp_docx_with_metadata_path = tmp_docx_with_metadata.name

            # Bước 2: Tạo mã QR
            qr_image = self.create_qr_code(public_key, self.current_user, metadata['sign_time'])
            
            # Bước 3: Thêm mã QR vào DOCX
            docx_path_with_qr_and_placeholder_signature = self.add_qr_to_docx(tmp_docx_with_metadata_path, qr_image)
            
            # Xóa file tạm thời ban đầu
            os.remove(tmp_docx_with_metadata_path)

            # Bước 4: Đọc lại file DOCX đã có QR để ký
            with open(docx_path_with_qr_and_placeholder_signature, 'rb') as f_to_sign:
                bytes_to_sign = f_to_sign.read()

            # Bước 5: Ký trên bytes của DOCX đã có QR
            signature = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64 = base64.b64encode(signature).decode('utf-8')
            
            # Bước 6: Cập nhật chữ ký thật vào document properties
            doc_final = Document(docx_path_with_qr_and_placeholder_signature)
            metadata_final = json.loads(doc_final.core_properties.comments or '{}')
            metadata_final['signature'] = signature_b64  # Cập nhật chữ ký thật
            doc_final.core_properties.comments = json.dumps(metadata_final)
            
            # Lưu file DOCX cuối cùng
            final_docx_path_for_upload = file_path.replace('.docx', '_signed_with_qr.docx')
            doc_final.save(final_docx_path_for_upload)
            
            # Xóa file tạm thời có QR và placeholder signature
            os.remove(docx_path_with_qr_and_placeholder_signature)

            # Upload file đã ký và có QR lên server
            with open(final_docx_path_for_upload, 'rb') as f:
                files = {'pdf': f}  # Giữ nguyên tên field để tương thích với server
                data = {
                    'signer_name': self.current_user,
                    'signature': signature_b64,
                    'timestamp': metadata['sign_time']
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-invoice/",
                    files=files,
                    data=data
                )
                response.raise_for_status()
            
            QMessageBox.information(
                self, "Thành công",
                f"Ký và upload DOCX thành công. File đã ký được lưu tại:\n{final_docx_path_for_upload}"
            )
            self.sign_path.clear()
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi ký và upload DOCX: {str(e)}")

    def refresh_invoice_list(self):
        """Làm mới danh sách hóa đơn"""
        if not self.current_user:
            return

        try:
            response = requests.get(
                f"{self.API_BASE_URL}/list-invoices/"
            )
            response.raise_for_status()
            
            invoices = response.json()
            self.invoice_list.clear()
            
            for invoice in invoices:
                item = QListWidgetItem(
                    f"Hóa đơn từ {invoice['signer_name']} - {invoice['timestamp']}"
                )
                item.setData(Qt.UserRole, invoice)
                self.invoice_list.addItem(item)
                
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách hóa đơn: {str(e)}")

    def download_and_verify_invoice(self, item):
        """Tải và xác thực hóa đơn DOCX"""
        invoice = item.data(Qt.UserRole)
        try:
            # Tải invoice
            response = requests.get(
                f"{self.API_BASE_URL}/download-invoice/{invoice['id']}"
            )
            response.raise_for_status()
            
            temp_path = os.path.join(os.path.expanduser("~"), "Downloads", f"invoice_{invoice['id']}_from_{invoice['signer_name']}.docx")
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            
            # Đọc file DOCX đã ký
            doc = Document(temp_path)
            metadata_str = doc.core_properties.comments
            if not metadata_str:
                QMessageBox.warning(self, "Lỗi", "File DOCX không có metadata.")
                return

            # Parse metadata
            metadata = json.loads(metadata_str)

            # Lấy thông tin người ký
            signer_seller = metadata.get('signer')
            signature_b64_seller = metadata.get('signature')
            sign_time_seller = metadata.get('sign_time')

            # Lấy certificate của người bán
            response_cert = requests.get(
                f"{self.API_BASE_URL}/get-certificate/",
                params={'username': signer_seller}
            )
            response_cert.raise_for_status()
            seller_certificate = response_cert.json()['certificate']

            # Xác minh certificate
            if not verify_certificate(seller_certificate):
                QMessageBox.warning(self, "Lỗi", "Certificate của người bán không hợp lệ")
                return

            # Lấy public key từ certificate đã xác minh
            public_key_seller = base64.b64decode(seller_certificate['payload']['public_key'])
            
            # Tạo lại bytes DOCX với placeholder để xác thực chữ ký
            doc_temp_seller = Document(temp_path)
            metadata_temp_seller = dict(metadata)
            metadata_temp_seller['signature'] = '__SIGNATURE_PLACEHOLDER__'
            doc_temp_seller.core_properties.comments = json.dumps(metadata_temp_seller)
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc_temp_seller.save(tmp_file.name)
                with open(tmp_file.name, 'rb') as f:
                    docx_bytes_with_placeholder_seller = f.read()
                os.remove(tmp_file.name)
            
            # Xác minh chữ ký
            signature_seller = base64.b64decode(signature_b64_seller)
            is_valid_seller_signature = ML_DSA_44.verify(public_key_seller, docx_bytes_with_placeholder_seller, signature_seller)

            # Hiển thị kết quả xác thực
            result_text = f"Kết quả xác thực hóa đơn:\n\n"
            result_text += f"Người bán: {signer_seller}\n"
            result_text += f"Thời gian ký: {sign_time_seller}\n"
            result_text += f"Chữ ký người bán: {'Hợp lệ' if is_valid_seller_signature else 'Không hợp lệ'}\n\n"

            if is_valid_seller_signature:
                result_text += "✅ Hóa đơn đã được xác thực thành công!"
            else:
                result_text += "❌ Hóa đơn không hợp lệ!"

            QMessageBox.information(self, "Kết quả xác thực", result_text)

        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi xác thực hóa đơn: {str(e)}")

    def get_key_paths(self, username=None):
        """Lấy đường dẫn đến file khóa"""
        # Sử dụng username được truyền vào hoặc current_user
        target_username = username or self.current_user
        if not target_username:
            return None, None
        
        private_key_path = os.path.join(self.keys_dir, f"{target_username}.private.pem")
        public_key_path = os.path.join(self.keys_dir, f"{target_username}.public.pem")
        return private_key_path, public_key_path

    def save_keys(self, public_key, private_key, password, username=None):
        """Lưu cặp khóa vào file"""
        try:
            # Đảm bảo thư mục keys tồn tại
            os.makedirs(self.keys_dir, exist_ok=True)
            
            # Sử dụng username được truyền vào hoặc current_user
            target_username = username or self.current_user
            if not target_username:
                print("Không có username để lưu khóa")  # Debug log
                return False
            
            private_key_path = os.path.join(self.keys_dir, f"{target_username}.private.pem")
            public_key_path = os.path.join(self.keys_dir, f"{target_username}.public.pem")
            
            print(f"Đang lưu khóa vào: {private_key_path} và {public_key_path}")  # Debug log

            # Lưu public key
            with open(public_key_path, 'w') as f:
                json.dump({
                    'public_key': base64.b64encode(public_key).decode('utf-8')
                }, f)

            # Mã hóa và lưu private key
            salt = get_random_bytes(16)
            iv = get_random_bytes(16)
            key = PBKDF2(password.encode(), salt, dkLen=32, count=10000)
            cipher = AES.new(key, AES.MODE_CBC, iv)
            
            # Pad private key
            padded_sk = private_key + b'\0' * (16 - len(private_key) % 16)
            encrypted_sk = cipher.encrypt(padded_sk)

            with open(private_key_path, 'w') as f:
                json.dump({
                    'salt': base64.b64encode(salt).decode('utf-8'),
                    'iv': base64.b64encode(iv).decode('utf-8'),
                    'encrypted_key': base64.b64encode(encrypted_sk).decode('utf-8')
                }, f)

            return True
        except Exception as e:
            logging.error(f"Lỗi khi lưu khóa: {str(e)}")
            print(f"Chi tiết lỗi: {str(e)}")  # Debug log
            return False

    def load_private_key(self, password):
        """Đọc và giải mã private key"""
        try:
            private_key_path, _ = self.get_key_paths()
            if not private_key_path or not os.path.exists(private_key_path):
                return None

            with open(private_key_path, 'r') as f:
                key_data = json.load(f)
                salt = base64.b64decode(key_data['salt'])
                iv = base64.b64decode(key_data['iv'])
                encrypted_key = base64.b64decode(key_data['encrypted_key'])

            key = PBKDF2(password.encode(), salt, dkLen=32, count=10000)
            cipher = AES.new(key, AES.MODE_CBC, iv)
            decrypted_key = cipher.decrypt(encrypted_key)
            return decrypted_key.rstrip(b'\0')

        except Exception as e:
            logging.error(f"Lỗi khi đọc private key: {str(e)}")
            return None

    def select_order_sign_file(self):
        """Chọn file Order DOCX để ký"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Chọn file Order DOCX", "", "Word Documents (*.docx)"
        )
        if file_path:
            self.order_sign_path.setText(file_path)

    def sign_and_upload_order_docx(self):
        """Bên mua ký và upload Order DOCX"""
        if not self.current_user:
            QMessageBox.warning(self, "Lỗi", "Vui lòng đăng nhập trước")
            return
        file_path = self.order_sign_path.text()
        if not file_path:
            QMessageBox.warning(self, "Lỗi", "Vui lòng chọn file Order DOCX")
            return
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu:", QLineEdit.Password
        )
        if not ok or not password:
            return
        try:
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key")
                return

            # Lấy public key từ server
            try:
                response = requests.get(
                    f"{self.API_BASE_URL}/get-public-key/",
                    params={'signer_name': self.current_user}
                )
                response.raise_for_status()
                public_key = base64.b64decode(response.json()['public_key'])
            except Exception as e:
                QMessageBox.critical(self, "Lỗi", f"Không thể lấy public key từ server: {str(e)}")
                return

            # Bước 1: Đọc DOCX gốc và tạo metadata
            doc = Document(file_path)
            
            # Tạo metadata cho document
            metadata = {
                'buyer': self.current_user,
                'signature': '__SIGNATURE_PLACEHOLDER__',  # Placeholder
                'sign_time': datetime.now().isoformat()
            }
            
            # Thêm metadata vào document properties
            doc.core_properties.author = self.current_user
            doc.core_properties.comments = json.dumps(metadata)

            # Lưu DOCX với metadata vào file tạm thời
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx_with_metadata:
                doc.save(tmp_docx_with_metadata.name)
                tmp_docx_with_metadata_path = tmp_docx_with_metadata.name

            # Bước 2: Tạo mã QR
            qr_image = self.create_qr_code(public_key, self.current_user, metadata['sign_time'])
            
            # Bước 3: Thêm mã QR vào DOCX
            docx_path_with_qr_and_placeholder_signature = self.add_qr_to_docx(tmp_docx_with_metadata_path, qr_image)
            
            # Xóa file tạm thời ban đầu
            os.remove(tmp_docx_with_metadata_path)

            # Bước 4: Đọc lại file DOCX đã có QR để ký
            with open(docx_path_with_qr_and_placeholder_signature, 'rb') as f_to_sign:
                bytes_to_sign = f_to_sign.read()

            # Bước 5: Ký trên bytes của DOCX đã có QR
            signature = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64 = base64.b64encode(signature).decode('utf-8')
            
            # Bước 6: Cập nhật chữ ký thật vào document properties
            doc_final = Document(docx_path_with_qr_and_placeholder_signature)
            metadata_final = json.loads(doc_final.core_properties.comments or '{}')
            metadata_final['signature'] = signature_b64  # Cập nhật chữ ký thật
            doc_final.core_properties.comments = json.dumps(metadata_final)
            
            # Lưu file DOCX cuối cùng
            final_docx_path_for_upload = file_path.replace('.docx', '_signed_with_qr.docx')
            doc_final.save(final_docx_path_for_upload)
            
            # Xóa file tạm thời có QR và placeholder signature
            os.remove(docx_path_with_qr_and_placeholder_signature)

            # Upload file đã ký và có QR lên server
            with open(final_docx_path_for_upload, 'rb') as f:
                files = {'pdf': f}  # Giữ nguyên tên field để tương thích với server
                data = {
                    'buyer_name': self.current_user,
                    'signature': signature_b64,
                    'timestamp': metadata['sign_time']
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-order/",
                    files=files,
                    data=data
                )
                response.raise_for_status()

            QMessageBox.information(
                self, "Thành công",
                f"Ký và upload Order DOCX thành công. File đã ký được lưu tại:\n{final_docx_path_for_upload}"
            )
            self.order_sign_path.clear()

        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi ký và upload Order DOCX: {str(e)}")

    def refresh_order_list(self):
        try:
            response = requests.get(f"{self.API_BASE_URL}/list-orders/")
            response.raise_for_status()
            orders = response.json()
            self.order_list.clear()
            for order in orders:
                item = QListWidgetItem(
                    f"Order từ {order['buyer_name']} - {order['timestamp']}"
                )
                item.setData(Qt.UserRole, order)
                self.order_list.addItem(item)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách order: {str(e)}")

    def refresh_my_order_list(self):
        try:
            response = requests.get(f"{self.API_BASE_URL}/list-orders/")
            response.raise_for_status()
            orders = response.json()
            self.my_order_list.clear()
            for order in orders:
                if order['buyer_name'] == self.current_user:
                    item = QListWidgetItem(
                        f"Order đã gửi - {order['timestamp']}"
                    )
                    item.setData(Qt.UserRole, order)
                    self.my_order_list.addItem(item)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi lấy danh sách order đã gửi: {str(e)}")

    def create_qr_code(self, public_key, signer_name, timestamp):
        """Tạo mã QR chứa thông tin khóa"""
        # Tạo dữ liệu cho mã QR
        qr_data = {
            "signer": signer_name,
            "timestamp": timestamp,
            "public_key": base64.b64encode(public_key).decode('utf-8')
        }
        qr_data_str = json.dumps(qr_data)
        
        # Tạo mã QR
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(qr_data_str)
        qr.make(fit=True)
        
        # Tạo hình ảnh mã QR
        qr_image = qr.make_image(fill_color="black", back_color="white")
        
        # Chuyển đổi hình ảnh thành bytes
        img_byte_arr = BytesIO()
        qr_image.save(img_byte_arr, format='PNG')
        return img_byte_arr.getvalue()

    def add_qr_to_pdf(self, pdf_path, qr_image_bytes):
        """Thêm mã QR vào trang cuối của PDF"""
        # Đọc PDF gốc (đã ký)
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Thêm tất cả các trang từ PDF gốc
        for page in reader.pages:
            writer.add_page(page)
        
        # Sao chép metadata từ PDF gốc sang writer mới
        if reader.metadata:
            writer.add_metadata(reader.metadata)

        # Tạo một trang mới với mã QR
        packet = BytesIO()
        c = canvas.Canvas(packet, pagesize=letter)
        
        # Vẽ mã QR
        # Lưu mã QR vào tệp tạm thời và sử dụng đường dẫn tệp
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_qr_file:
            temp_qr_file.write(qr_image_bytes)
            temp_qr_path = temp_qr_file.name

        c.drawImage(temp_qr_path, 50, 50, width=200, height=200)
        c.save()
        
        # Xóa tệp QR tạm thời
        os.remove(temp_qr_path)
        
        # Thêm trang mới (chứa QR) vào PDF
        packet.seek(0)
        new_pdf = PdfReader(packet)
        writer.add_page(new_pdf.pages[0])
        
        # Lưu PDF mới
        output_path = pdf_path.replace('.pdf', '_with_qr.pdf')
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return output_path

    def show_register_dialog(self):
        """Hiển thị dialog đăng ký"""
        dialog = RegisterDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            self.register_user(data)

    def register_user(self, data):
        """Xử lý đăng ký người dùng"""
        username = data['username']
        password = data['password']
        confirm_password = data['confirm_password']
        role = data['role']
        
        if not username or not password or not role:
            QMessageBox.warning(self, "Lỗi", "Vui lòng điền đầy đủ thông tin và chọn vai trò")
            return
        
        if password != confirm_password:
            QMessageBox.warning(self, "Lỗi", "Mật khẩu xác nhận không khớp")
            return
        
        try:
            # Kiểm tra username đã tồn tại chưa
            response = requests.get(
                f"{self.API_BASE_URL}/check-user",
                params={"username": username}
            )
            response_data = response.json()
            
            if response_data.get('message') == 'User exists':
                QMessageBox.warning(self, "Lỗi", "Tên đăng nhập đã tồn tại")
                return
            
            # Tạo cặp khóa mới
            pk, sk = ML_DSA_44.keygen()
            
            # Tạo payload để chứng minh sở hữu private key
            payload = {
                "username": username,
                "timestamp": datetime.utcnow().isoformat(),
                "role": role
            }
            
            # Ký payload bằng private key
            signature = ML_DSA_44.sign(sk, json.dumps(payload, sort_keys=True).encode())
            signature_b64 = base64.b64encode(signature).decode('utf-8')
            
            # Lưu khóa vào file local
            if not self.save_keys(pk, sk, password, username):
                raise Exception("Không thể lưu cặp khóa vào file")
            
            # Gửi request đăng ký
            response = requests.post(
                f"{self.API_BASE_URL}/register/",
                json={
                    'public_key': base64.b64encode(pk).decode('utf-8'),
                    'payload': payload,
                    'signature': signature_b64
                }
            )
            response.raise_for_status()
            response_data = response.json()

            # Verify CA public key trước khi lưu
            ca_public_key = response_data.get('ca_public_key')
            ca_public_key_hash = response_data.get('ca_public_key_hash')
            
            if not ca_public_key or not ca_public_key_hash:
                raise Exception("Không nhận được CA public key hoặc hash")
            
            # Verify hash của CA public key
            ca_public_key_bytes = base64.b64decode(ca_public_key)
            if hashlib.sha256(ca_public_key_bytes).digest() != base64.b64decode(ca_public_key_hash):
                raise Exception("CA public key hash không hợp lệ")
            
            # Nếu verify thành công, lưu CA info
            ca_info = {
                'ca_public_key': ca_public_key,
                'ca_public_key_hash': ca_public_key_hash
            }
            ca_info_path = os.path.join(self.keys_dir, 'ca_info.json')
            with open(ca_info_path, 'w') as f:
                json.dump(ca_info, f)

            # Lưu certificate nếu có
            
            QMessageBox.information(
                self, "Thành công",
                "Đăng ký thành công! Bạn có thể đăng nhập ngay bây giờ."
            )
            self.username_input.setText(username)

        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi đăng ký: {str(e)}")

    def sign_downloaded_order_as_invoice(self, order_docx_path, buyer_name):
        """Người bán ký Order đã tải về và tạo Invoice với mã QR."""
        password, ok = QInputDialog.getText(
            self, "Nhập mật khẩu", "Mật khẩu của bạn:", QLineEdit.Password
        )
        if not ok or not password:
            return
        
        try:
            private_key = self.load_private_key(password)
            if not private_key:
                QMessageBox.critical(self, "Lỗi", "Không thể đọc private key của bạn.")
                return

            # Lấy certificate của người bán từ server
            try:
                response = requests.get(
                    f"{self.API_BASE_URL}/get-certificate/",
                    params={'username': self.current_user}
                )
                response.raise_for_status()
                seller_certificate = response.json()['certificate']
                
                # Verify certificate
                if not verify_certificate(seller_certificate):
                    QMessageBox.critical(self, "Lỗi", "Certificate của bạn không hợp lệ")
                    return
                    
                # Lấy public key từ certificate đã verify
                public_key_seller = base64.b64decode(seller_certificate['payload']['public_key'])
                
            except Exception as e:
                QMessageBox.critical(self, "Lỗi", f"Không thể lấy certificate của bạn từ server: {str(e)}")
                return

            # Bước 1: Đọc Order DOCX đã tải về và giữ lại metadata hiện có (chữ ký của người mua)
            doc_order = Document(order_docx_path)
            metadata_order_str = doc_order.core_properties.comments
            metadata_order = json.loads(metadata_order_str) if metadata_order_str else {}
            
            # Tạo metadata cho Invoice
            metadata_invoice = {
                'buyer': buyer_name,  # Giữ lại thông tin người mua
                'buyer_signature': metadata_order.get('signature'),  # Giữ lại chữ ký người mua
                'buyer_sign_time': metadata_order.get('sign_time'),  # Giữ lại thời gian ký của người mua
                'seller': self.current_user,  # Thêm thông tin người bán
                'seller_signature': '__SIGNATURE_PLACEHOLDER_SELLER__',  # Placeholder chữ ký người bán
                'seller_sign_time': datetime.now().isoformat()  # Thời gian ký của người bán
            }
            
            # Cập nhật document properties
            doc_order.core_properties.author = self.current_user
            doc_order.core_properties.comments = json.dumps(metadata_invoice)

            # Lưu DOCX này vào một tệp tạm thời để thêm mã QR
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx_with_metadata_invoice:
                doc_order.save(tmp_docx_with_metadata_invoice.name)
                tmp_docx_with_metadata_invoice_path = tmp_docx_with_metadata_invoice.name
            
            # Bước 2: Tạo mã QR với thông tin của người bán
            qr_image = self.create_qr_code(public_key_seller, self.current_user, metadata_invoice['seller_sign_time'])
            
            # Bước 3: Thêm mã QR vào DOCX tạm thời
            docx_path_with_qr_and_placeholder_signature_seller = self.add_qr_to_docx(tmp_docx_with_metadata_invoice_path, qr_image)
            
            # Xóa file tạm thời ban đầu
            os.remove(tmp_docx_with_metadata_invoice_path)

            # Bước 4: Đọc lại file DOCX đã có QR để ký
            with open(docx_path_with_qr_and_placeholder_signature_seller, 'rb') as f_to_sign:
                bytes_to_sign = f_to_sign.read()

            # Bước 5: Ký trên bytes của DOCX đã có QR
            signature_seller = ML_DSA_44.sign(private_key, bytes_to_sign)
            signature_b64_seller = base64.b64encode(signature_seller).decode('utf-8')
            
            # Bước 6: Cập nhật chữ ký thật vào document properties
            doc_final_invoice = Document(docx_path_with_qr_and_placeholder_signature_seller)
            metadata_final_invoice = json.loads(doc_final_invoice.core_properties.comments or '{}')
            metadata_final_invoice['seller_signature'] = signature_b64_seller  # Cập nhật chữ ký thật
            doc_final_invoice.core_properties.comments = json.dumps(metadata_final_invoice)
            
            # Lưu file DOCX cuối cùng
            final_invoice_docx_path = os.path.join(os.path.expanduser("~"), "Downloads", f"invoice_{buyer_name}_from_{self.current_user}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc_final_invoice.save(final_invoice_docx_path)
            
            # Xóa file tạm thời có QR và placeholder signature
            os.remove(docx_path_with_qr_and_placeholder_signature_seller)

            # Upload Invoice DOCX đã ký (và có QR) lên server
            with open(final_invoice_docx_path, 'rb') as f:
                files = {'pdf': f}  # Giữ nguyên tên field để tương thích với server
                data = {
                    'signer_name': self.current_user,
                    'signature': signature_b64_seller,
                    'timestamp': metadata_invoice['seller_sign_time']
                }
                response = requests.post(
                    f"{self.API_BASE_URL}/upload-invoice/",
                    files=files,
                    data=data
                )
                response.raise_for_status()

            QMessageBox.information(
                self, "Thành công",
                f"Đã ký Order và tạo Hóa đơn thành công. Hóa đơn được lưu tại:\n{final_invoice_docx_path}"
            )
            self.refresh_order_list()
            
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi ký Order và tạo Hóa đơn: {str(e)}")

    def add_qr_to_docx(self, docx_path, qr_image_bytes):
        """Thêm mã QR vào cuối file Word DOCX"""
        try:
            # Đọc file DOCX gốc
            doc = Document(docx_path)
            
            # Thêm một trang mới (section break) để chứa QR code
            doc.add_page_break()
            
            # Thêm tiêu đề cho phần QR code
            qr_heading = doc.add_heading('Digital Signature QR Code', level=1)
            qr_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm đoạn văn bản mô tả
            description = doc.add_paragraph('This QR code contains the digital signature information for this document.')
            description.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Lưu QR code vào file tạm thời
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_qr_file:
                temp_qr_file.write(qr_image_bytes)
                temp_qr_path = temp_qr_file.name
            
            # Thêm QR code vào document
            qr_paragraph = doc.add_paragraph()
            qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr_run = qr_paragraph.add_run()
            qr_run.add_picture(temp_qr_path, width=Inches(2.5), height=Inches(2.5))
            
            # Xóa file QR tạm thời
            os.remove(temp_qr_path)
            
            # Thêm thông tin bổ sung
            info_paragraph = doc.add_paragraph()
            info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_paragraph.add_run('Scan this QR code to verify the digital signature of this document.')
            
            # Lưu file DOCX mới
            output_path = docx_path.replace('.docx', '_with_qr.docx')
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"Error adding QR to DOCX: {str(e)}")
            raise

    def download_and_verify_order(self, item):
        """Tải và xác thực order DOCX"""
        order = item.data(Qt.UserRole)
        try:
            response = requests.get(
                f"{self.API_BASE_URL}/download-order/{order['id']}"
            )
            response.raise_for_status()
            
            temp_path = os.path.join(os.path.expanduser("~"), "Downloads", f"order_{order['id']}_from_{order['buyer_name']}.docx")
            with open(temp_path, 'wb') as f:
                f.write(response.content)
            
            # Đọc file DOCX đã ký
            doc = Document(temp_path)
            metadata_str = doc.core_properties.comments
            if not metadata_str:
                QMessageBox.warning(self, "Lỗi", "File DOCX không có metadata.")
                return

            # Parse metadata
            metadata = json.loads(metadata_str)
            
            signature_b64 = metadata.get('signature')
            buyer = metadata.get('buyer')
            sign_time = metadata.get('sign_time')
            
            if not signature_b64 or not buyer:
                QMessageBox.warning(self, "Lỗi", "File DOCX không có chữ ký hợp lệ")
                return
            
            # Lấy certificate của người mua
            response_cert = requests.get(
                f"{self.API_BASE_URL}/get-certificate/",
                params={'username': buyer}
            )
            response_cert.raise_for_status()
            buyer_certificate = response_cert.json()['certificate']

            # Xác minh certificate
            if not verify_certificate(buyer_certificate):
                QMessageBox.warning(self, "Lỗi", "Certificate của người mua không hợp lệ")
                return

            # Lấy public key từ certificate đã xác minh
            public_key_buyer = base64.b64decode(buyer_certificate['payload']['public_key'])
            
            # Tạo lại bytes DOCX với placeholder để xác thực chữ ký
            doc_temp = Document(temp_path)
            metadata_temp = dict(metadata)
            metadata_temp['signature'] = '__SIGNATURE_PLACEHOLDER__'
            doc_temp.core_properties.comments = json.dumps(metadata_temp)
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc_temp.save(tmp_file.name)
                with open(tmp_file.name, 'rb') as f:
                    docx_bytes_with_placeholder = f.read()
                os.remove(tmp_file.name)
            
            signature = base64.b64decode(signature_b64)
            is_valid_signature = ML_DSA_44.verify(public_key_buyer, docx_bytes_with_placeholder, signature)
            
            if is_valid_signature:
                reply = QMessageBox.information(
                    self, "Xác thực Order thành công!",
                    f"Order từ {buyer} - Thời gian: {sign_time}\n"
                    f"Chữ ký của người mua hợp lệ.\n\nBạn có muốn ký và tạo Hóa đơn (Invoice) từ Order này không?",
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
                )
                if reply == QMessageBox.Yes:
                    self.sign_downloaded_order_as_invoice(temp_path, buyer)
                elif reply == QMessageBox.No:
                    QMessageBox.information(self, "Thông báo", "Order đã được tải về và xác thực.")
            else:
                QMessageBox.warning(self, "Lỗi", "Xác thực thất bại: Chữ ký của người mua không hợp lệ.")

        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi khi tải và xác thực order: {str(e)}")

class RegisterDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Đăng ký tài khoản")
        self.setModal(True)
        self.init_ui()

    def init_ui(self):
        layout = QFormLayout(self)
        
        # Thêm các trường nhập liệu
        self.username_input = QLineEdit()
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        self.confirm_password_input = QLineEdit()
        self.confirm_password_input.setEchoMode(QLineEdit.Password)

        self.role_selector = QComboBox()
        self.role_selector.addItem("Người mua", "buyer")
        self.role_selector.addItem("Người bán", "seller")
        
        layout.addRow("Tên đăng nhập:", self.username_input)
        layout.addRow("Mật khẩu:", self.password_input)
        layout.addRow("Xác nhận mật khẩu:", self.confirm_password_input)
        layout.addRow("Vai trò:", self.role_selector)
        
        # Thêm nút đăng ký
        register_btn = QPushButton("Đăng ký")
        register_btn.clicked.connect(self.accept)
        layout.addRow("", register_btn)

    def get_data(self):
        return {
            'username': self.username_input.text().strip(),
            'password': self.password_input.text(),
            'confirm_password': self.confirm_password_input.text(),
            'role': self.role_selector.currentData()
        }

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = PDFVerifier()
    window.show()
    sys.exit(app.exec_()) 